# main_window.py
import sys
import json
import logging
from datetime import datetime
from PyQt6.QtCore import Qt, QTimer
from PyQt6.QtGui import QKeySequence, QShortcut, QColor, QPixmap, QIcon
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLineEdit, QTableWidgetItem,
    QFileDialog, QLabel, QMessageBox, QSplitter, QHeaderView,
    QMenu, QAbstractItemView, QDialog, QTreeWidget, QTreeWidgetItem, QFrame, QProgressBar
)
from custom_table_widget import CustomTableWidget
from app_constants import ColumnIndex

import database
import config
import themes
import utils
from models import Shipment, ShipmentItem, Box, GroupShipment, ShipmentProperties
from dialogs import GoogleSheetsImportDialog, SettingsDialog, DestinationDialog, RenameDialog, ShipmentPropertiesDialog, ArchiveDialog, QuantityEditDialog, GroupShipmentPropertiesDialog, GoogleSheetsUpdateDialog
from db_settings_dialog import DatabaseSettingsDialog
from shipment_manager import ShipmentManager
from shipment_operations import ShipmentOperations
from ui_updater import UIUpdater
from editing_delegate import QuantityEditDelegate
from pathlib import Path
from shipment_check_dialog import ShipmentCheckDialog

# Импорт исключения из db_connection (поддерживает psycopg2 и psycopg3)
from db_connection import psycopg, psycopg2
from image_check_box import ImageCheckBox

# Добавляем импорт для работы с Word
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from shipment_controller import ShipmentController
from data_controller import DataController
from async_operations import AsyncOperationsManager
# MainController импортируется лениво в _lazy_init_controllers() для избежания циклических импортов


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.logger = logging.getLogger(__name__)
        self.logger.info("Инициализация главного окна")
        
        # Получаем версию приложения
        try:
            from version import get_version_string
            version = get_version_string()
            self.setWindowTitle(f"WB Packer — Сборка поставок (v{version})")
        except Exception:
            self.setWindowTitle("WB Packer — Сборка поставок")
        
        self.setWindowIcon(QIcon(str(config.get_resource_path(Path("Res") / "icon.ico"))))
        self.resize(1300, 800)
        
        # Устанавливаем политику фокуса, чтобы окно могло получать фокус
        self.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        
        # Инициализация статус-бара
        self.statusBar().showMessage("Готово")
        self.statusBar().setSizeGripEnabled(True)
        
        # Прогресс-бар в статус-баре
        self.status_progress_bar = QProgressBar()
        self.status_progress_bar.setMaximumWidth(200)
        self.status_progress_bar.setVisible(False)
        self.statusBar().addPermanentWidget(self.status_progress_bar)
        
        # Инициализация переменных
        self.saving = False
        self.updating_ui = False
        self.conflict_items = set()
        self.last_completed_state = {}
        self.initialization_complete = False  # Флаг завершения инициализации
        
        # Настройки пользователя
        self.current_user = None
        self.font_size = config.DEFAULT_FONT_SIZE
        self.label_font_size = config.DEFAULT_LABEL_FONT_SIZE
        self.current_theme = config.DEFAULT_THEME
        self.ok_sound = "ok.wav"
        self.error_sound = "error.wav"
        self.tone_sound = False
        self.sound_volume = 100
        self.shipment_columns_width = ""
        self.box_columns_width = ""
        self.main_splitter_sizes = ""
        self.window_width = 1300
        self.window_height = 800
        self.button_colors = {}  # Словарь для хранения цветов кнопок
        self.shipment_locking_enabled = True  # По умолчанию блокировка включена
        
        # Таймер для отложенного сохранения настроек
        self.save_timer = QTimer()
        self.save_timer.setSingleShot(True)
        self.save_timer.timeout.connect(self.save_user_settings)
        
        # Флаг для отслеживания первой загрузки
        self.first_load = True
        
        # Данные
        self.current_shipment = None
        self.shipments = {}
        self.group_shipments = {}
        
        # Инициализация контроллеров (отложенно)
        self.shipment_controller = None

        # Инициализация менеджеров (отложенно)
        self.shipment_manager = None
        self.shipment_operations = None
        self.ui_updater = None
        self.data_controller = None
        self.async_manager = None
        self.mvc_controller = None
        
        # Инициализация
        try:
            # Только базовая инициализация - остальное сделаем после показа окна
            self.init_database()
            self.init_ui()
            self.setup_shortcuts()
            self.initialize_user()

            # Отложенная загрузка данных и применение настроек
            # Вызывается из main.py после показа окна
            # QTimer.singleShot(100, self.deferred_initialization)
        except Exception as e:
            self.logger.error(f"Ошибка при инициализации главного окна: {e}", exc_info=True)
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при запуске приложения:\n{e}")
        
        # if not self.shipments and not self.group_shipments:
        #     # Статус скрыт, сообщения не отображаются
        #     # self.statusBar().showMessage("Нажмите «+ Поставка (F1)», чтобы начать")
        

    def _lazy_init_controllers(self):
        """Ленивая инициализация контроллеров"""
        if self.shipment_controller is None:
            from shipment_controller import ShipmentController
            self.shipment_controller = ShipmentController(self)
    
    def _lazy_init_managers(self):
        """Ленивая инициализация менеджеров"""
        if self.shipment_manager is None:
            from shipment_manager import ShipmentManager
            self.shipment_manager = ShipmentManager(self)
            
        if self.shipment_operations is None:
            from shipment_operations import ShipmentOperations
            self.shipment_operations = ShipmentOperations(self)
            
        if self.ui_updater is None:
            from ui_updater import UIUpdater
            self.ui_updater = UIUpdater(self)
            
        if self.data_controller is None:
            from data_controller import DataController
            self.data_controller = DataController(self)
            
        if self.async_manager is None:
            from async_operations import AsyncOperationsManager
            self.async_manager = AsyncOperationsManager()
            
        if self.mvc_controller is None:
            from mvc_controller import MainController
            self.mvc_controller = MainController(self)
    
    def deferred_initialization(self):
        """Отложенная инициализация после показа окна"""
        try:
            self.logger.info("Завершение загрузки состояния окна")

            # Ленивая инициализация контроллеров и менеджеров
            self._lazy_init_controllers()
            self._lazy_init_managers()

            # Привязываем события после инициализации MVC контроллера
            self.mvc_controller.bind_events()

            # Сначала загружаем настройки пользователя (тему, шрифты и т.д.)
            self._loading_settings = True
            self.load_user_settings()

            # Загружаем данные ИЗ БД ДО применения темы
            # Это гарантирует, что дерево поставок будет заполнено данными
            self.load_all_data()

            # Теперь применяем настройки темы и UI ПОСЛЕ загрузки данных
            # Это обеспечит корректное отображение с правильной темой
            self.apply_settings()

            # Принудительно обновляем стили таблиц ПОСЛЕ заполнения
            self._update_table_styles()

            # Восстанавливаем ширину столбцов таблиц
            self.restore_table_columns_width()

            # Применяем настройки видимости к чекбоксам и столбцам (ПОСЛЕ apply_settings)
            self._apply_visibility_settings_to_checkboxes()

            # Восстанавливаем состояние окна
            self.load_window_state()

            # Завершаем загрузку настроек
            self._loading_settings = False

            # Помечаем инициализацию как завершенную
            self.initialization_complete = True

            self.logger.info("Отложенная инициализация завершена")

            # if not self.shipments and not self.group_shipments:
            #     # Статус скрыт, сообщения не отображаются
            #     # self.statusBar().showMessage("Нажмите «+ Поставка (F1)», чтобы начать")

        except Exception as e:
            self.logger.error(f"Ошибка при отложенной инициализации: {e}", exc_info=True)
            QMessageBox.warning(self, "Ошибка", f"Не удалось завершить инициализацию: {e}")
    
    def cleanup_old_sessions(self):
        """Очистка старых сессий пользователей"""
        try:
            from database import cleanup_old_sessions
            cleanup_old_sessions()
        except Exception as e:
            self.logger.error(f"Ошибка при очистке старых сессий пользователей: {e}", exc_info=True)
    
    def update_user_activity(self):
        """Обновление кактивности пользователя в текущей поставке"""
        if (self.current_shipment and
            hasattr(self, 'current_user') and
            self.current_user):
            try:
                from database import update_user_session
                update_user_session(self.current_shipment.destination_name, self.current_user)

                # Ткакже добавляем пользователя в список активных пользователей текущей поставки
                self.current_shipment.add_active_user(self.current_user)
            except Exception as e:
                self.logger.error(f"Ошибка при обновлении сессии пользователя: {e}", exc_info=True)

    def init_database(self):
        """Инициализация базы данных"""
        try:
            # Проверяем, была ли уже инициализирована база данных
            # Это предотвращает двойную инициализацию
            if not hasattr(database, '_initialized') or not database._initialized:
                database.init_db()
                database._initialized = True  # Помечаем как инициализированную
            
            # Обновляем индикатор базы данных
            self._update_db_indicator()
        except Exception as e:
            # Универсальная обработка ошибок для psycopg2 и psycopg3
            self.logger.error(f"Ошибка подключения к базе данных: {e}", exc_info=True)
            error_str = str(e).lower()
            
            # Проверяем тип ошибки
            is_db_error = False
            if psycopg is not None:
                is_db_error = isinstance(e, (psycopg.OperationalError, psycopg.DatabaseError))
            elif psycopg2 is not None:
                is_db_error = isinstance(e, (psycopg2.OperationalError, psycopg2.DatabaseError))
            
            if is_db_error or "connection refused" in error_str or "failed:" in error_str or "timeout" in error_str or "byte" in error_str or "could not connect" in error_str:
                QMessageBox.critical(self, "Критическая ошибка",
                                   f"Не удалось подключиться к PostgreSQL базе данных.\n\n"
                                   f"Проверьте настройки подключения:\n"
                                   f"  - Сервер: {config.get_postgresql_host()}\n"
                                   f"  - Порт: {config.POSTGRESQL_PORT}\n"
                                   f"  - База данных: {config.POSTGRESQL_DATABASE}\n"
                                   f"  - Пользователь: {config.POSTGRESQL_USER}\n\n"
                                   f"Убедитесь, что сервер PostgreSQL запущен и доступен.")
                sys.exit(1)
            else:
                QMessageBox.critical(self, "Ошибка БД", f"Не удалось инициализировать базу данных:\n{e}")
                sys.exit(1)

    def _init_menu_bar(self):
        """Инициализация меню "Сервис" и его пунктов"""
        menubar = self.menuBar()
        service_menu = menubar.addMenu("Сервис")

        self.user_action = service_menu.addAction("👤 Управление пользователями")
        self.user_action.triggered.connect(self.manage_users)

        settings_action = service_menu.addAction("⚙️ Настройки")
        settings_action.triggered.connect(self.open_settings)

        archive_action = service_menu.addAction("📦 Архив поставок")
        archive_action.triggered.connect(self.open_archive)

        service_menu.addSeparator()

        about_action = service_menu.addAction("ℹ️ О программе")
        about_action.triggered.connect(self.show_about)

        # Добавляем индикатор базы данных в правую часть меню
        self._init_db_indicator()

    def _init_db_indicator(self):
        """Инициализация индикатора подключения базы данных"""
        from PyQt6.QtWidgets import QWidget, QHBoxLayout, QLabel
        from PyQt6.QtCore import Qt
        
        # Создаём виджет-контейнер для индикатора
        indicator_widget = QWidget()
        indicator_layout = QHBoxLayout(indicator_widget)
        indicator_layout.setContentsMargins(10, 2, 10, 2)
        indicator_layout.setSpacing(8)
        
        # Создаём цветную точку
        self.db_dot_label = QLabel("●")
        self.db_dot_label.setStyleSheet("font-size: 14px;")
        self.db_dot_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # Создаём надпись БД
        db_text_label = QLabel("БД")
        db_text_label.setStyleSheet("font-size: 12px; font-weight: bold;")
        
        # Создаём метку для IP адреса БД
        self.db_host_label = QLabel("")
        self.db_host_label.setStyleSheet("font-size: 12px; color: #888;")
        
        # Создаём разделитель
        separator_label = QLabel("•")
        separator_label.setStyleSheet("font-size: 12px; color: #888;")
        
        # Создаём метку для даты
        self.date_label = QLabel("")
        self.date_label.setStyleSheet("font-size: 12px; color: #888;")
        
        indicator_layout.addWidget(self.db_dot_label)
        indicator_layout.addWidget(db_text_label)
        indicator_layout.addWidget(self.db_host_label)
        indicator_layout.addWidget(separator_label)
        indicator_layout.addWidget(self.date_label)
        indicator_layout.addStretch()
        
        # Добавляем индикатор в меню бар
        self.menuBar().setCornerWidget(indicator_widget, Qt.Corner.TopRightCorner)
        
        # Обновляем цвет индикатора и данные
        self._update_db_indicator()
    
    def _update_db_indicator(self):
        """Обновление цвета индикатора базы данных и информации"""
        if hasattr(self, 'db_dot_label'):
            db_type = database.get_db_type()
            if db_type == "postgresql":
                # Зелёная точка для PostgreSQL
                self.db_dot_label.setStyleSheet("font-size: 14px; color: #4CAF50;")
                # Получаем хост PostgreSQL
                host = config.get_postgresql_host()
                self.db_host_label.setText(f"{host}")
            else:
                # Голубая точка для SQLite
                self.db_dot_label.setStyleSheet("font-size: 14px; color: #2196F3;")
                self.db_host_label.setText("SQLite")
            
            # Обновляем дату
            self._update_date_label()
    
    def _update_date_label(self):
        """Обновление метки с текущей датой"""
        if hasattr(self, 'date_label'):
            from datetime import datetime
            now = datetime.now()
            # Формат: "13 марта 2026"
            date_str = now.strftime("%d %B %Y").replace(now.strftime("%B"), 
                {
                    1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
                    5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
                    9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
                }[now.month])
            self.date_label.setText(date_str)

    def _init_main_layout(self):
        """Создание центрального виджета и основного layout"""
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(5, 5, 5, 5)
        main_layout.setSpacing(8)
        return main_layout

    def _init_top_buttons_panel(self):
        """Создание верхней панели с кнопками действий"""
        top_layout = QHBoxLayout()
        buttons_layout = QHBoxLayout()

        self.new_shipment_btn = QPushButton("+ Поставка")
        self.new_shipment_gh_btn = QPushButton("+ Поставка GH")
        self.new_box_btn = QPushButton("+ Коробка")
        self.shipment_check_btn = QPushButton("Проверка")
        self.print_labels_btn = QPushButton("Этикетки")
        self.moysklad_sync_btn = QPushButton("МойСклад")
        self.check_stock_btn = QPushButton("Остаток")

        # Кнопки подключены в mvc_controller.py через контроллер
        # self.new_shipment_btn.clicked.connect(self.start_new_shipment)
        # self.new_box_btn.clicked.connect(self.new_box)
        # self.shipment_check_btn.clicked.connect(self.start_shipment_check)
        # self.print_labels_btn.clicked.connect(...)
        # Эти кнопки подключаем напрямую:
        self.moysklad_sync_btn.clicked.connect(self.sync_moysklad_stocks)
        self.check_stock_btn.clicked.connect(self.open_check_stock_dialog)
        self.new_shipment_gh_btn.clicked.connect(self.import_shipment_from_google_sheets)

        buttons_layout.addWidget(self.new_shipment_btn)
        buttons_layout.addWidget(self.new_shipment_gh_btn)
        buttons_layout.addWidget(self.new_box_btn)
        buttons_layout.addWidget(self.shipment_check_btn)
        buttons_layout.addWidget(self.print_labels_btn)
        buttons_layout.addWidget(self.moysklad_sync_btn)
        buttons_layout.addWidget(self.check_stock_btn)
        self.update_moysklad_button_visibility()

        # Применяем стили к кнопкам после создания
        self.apply_button_styles()

        return top_layout, buttons_layout

    def apply_button_styles(self):
        """Применяет цветные стили к кнопкам - всегда цветные"""
        # Используем новый метод с фиксированными цветами
        self._apply_fixed_colorful_button_styles()

        # Принудительно обновляем кнопки
        buttons = [self.new_shipment_btn, self.new_shipment_gh_btn, self.new_box_btn,
                   self.shipment_check_btn, self.print_labels_btn, self.moysklad_sync_btn,
                   self.check_stock_btn]
        for btn in buttons:
            btn.style().unpolish(btn)
            btn.style().polish(btn)

    def apply_toggle_styles(self):
        """Применяет стили к чекбоксам с изображениями"""
        if not hasattr(self, 'article_display_checkbox'):
            return

        try:
            theme = themes.THEMES.get(self.current_theme, themes.THEMES["Светлая"])
            text_color = theme["window_text"]
            bg_color = theme["window_bg"]
        except Exception as e:
            self.logger.error(f"Ошибка получения темы в apply_toggle_styles: {e}")
            # Используем цвета по умолчанию
            text_color = QColor(40, 40, 40)  # Тёмный текст для светлой темы
            bg_color = QColor(245, 245, 247)

        checkboxes = []
        if hasattr(self, 'article_display_checkbox'):
            checkboxes.append(self.article_display_checkbox)
        if hasattr(self, 'name_display_checkbox'):
            checkboxes.append(self.name_display_checkbox)
        if hasattr(self, 'total_qty_display_checkbox'):
            checkboxes.append(self.total_qty_display_checkbox)
        if hasattr(self, 'stock_display_checkbox'):
            checkboxes.append(self.stock_display_checkbox)
        if hasattr(self, 'hide_completed_checkbox'):
            checkboxes.append(self.hide_completed_checkbox)

        for checkbox in checkboxes:
            if checkbox:
                # Устанавливаем цвет текста через палитру - это работает надежнее стилей
                from PyQt6.QtGui import QPalette
                palette = checkbox.palette()
                palette.setColor(checkbox.foregroundRole(), text_color)
                palette.setColor(checkbox.backgroundRole(), bg_color)
                palette.setColor(QPalette.ColorRole.WindowText, text_color)
                palette.setColor(QPalette.ColorRole.Window, bg_color)
                palette.setColor(QPalette.ColorRole.Text, text_color)
                palette.setColor(QPalette.ColorRole.Button, bg_color)
                palette.setColor(QPalette.ColorRole.ButtonText, text_color)
                checkbox.setPalette(palette)
                
                # Устанавливаем атрибут для использования палитры
                checkbox.setAttribute(Qt.WidgetAttribute.WA_SetPalette, True)
                
                # Используем метод setTextColor из ImageCheckBox
                if hasattr(checkbox, 'setTextColor'):
                    checkbox.setTextColor(text_color)
                
                # Ткакже устанавливаем цвет напрямую через стиль
                checkbox.setStyleSheet(f"""
                    QCheckBox {{
                        color: {text_color.name()};
                        background-color: {bg_color.name()};
                        font-size: 14px;
                        spacing: 12px;
                    }}
                    QCheckBox::indicator {{
                        width: 0;
                        height: 0;
                    }}
                """)
                
                # Принудительно обновляем иконку после применения стиля
                if hasattr(checkbox, '_update_icon'):
                    checkbox._update_icon()

                # Принудительно обновляем виджет
                checkbox.update()
                checkbox.repaint()

    def _update_table_styles(self):
        """Принудительно обновляет стили таблиц после применения темы"""
        try:
            theme = themes.THEMES.get(self.current_theme, themes.THEMES["Светлая"])

            # Получаем цвет для чередования строк
            highlight_color = theme["highlight"].name()
            table_bg_color = theme["table_bg"].name()

            # Обновляем стили для shipment_table
            if hasattr(self, 'shipment_table') and self.shipment_table:
                self.shipment_table.setStyleSheet(f"""
                    QTableWidget {{
                        background-color: {table_bg_color};
                        alternate-background-color: {highlight_color};
                        gridline-color: transparent;
                        border: none;
                    }}
                """)
                self.shipment_table.setAlternatingRowColors(True)
                # Принудительно обновляем таблицу
                self.shipment_table.viewport().update()
                self.shipment_table.repaint()

            # Обновляем стили для current_box_table
            if hasattr(self, 'current_box_table') and self.current_box_table:
                self.current_box_table.setStyleSheet(f"""
                    QTableWidget {{
                        background-color: {table_bg_color};
                        alternate-background-color: {highlight_color};
                        gridline-color: transparent;
                        border: none;
                    }}
                """)
                self.current_box_table.setAlternatingRowColors(True)
                # Принудительно обновляем таблицу
                self.current_box_table.viewport().update()
                self.current_box_table.repaint()
                
        except Exception as e:
            self.logger.error(f"Ошибка обновления стилей таблиц: {e}")

    def _style_top_button(self, button, bg_color, hover_color):
        """Стилизация кнопки верхней панели"""
        button.setStyleSheet(f"""
            QPushButton {{
                background-color: {bg_color};
                color: white;
                border: none;
                border-radius: 4px;
                padding: 4px 10px;
                font-weight: bold;
                font-size: 11px;
                min-width: 70px;
                max-width: 90px;
            }}
            QPushButton:hover {{
                background-color: {hover_color};
            }}
            QPushButton:pressed {{
                background-color: {bg_color};
                padding-top: 6px;
                padding-bottom: 2px;
            }}
            QPushButton:disabled {{
                background-color: #cccccc;
                color: #666666;
            }}
        """)

    def _init_display_controls(self, top_layout, buttons_layout):
        """Создание чекбоксов для управления отображением столбцов и строк"""
        theme_layout = QHBoxLayout()
        theme_layout.addStretch()

        # Инициализируем переменные видимости значениями по умолчанию
        self.article_column_visible = True
        self.name_column_visible = False
        self.total_qty_column_visible = True
        self.stock_column_visible = True
        self.hide_completed_items_setting = False

        self.article_display_checkbox = ImageCheckBox("Артикул")
        self.article_display_checkbox.setChecked(self.article_column_visible)
        self.article_display_checkbox.stateChanged.connect(self.toggle_article_column)
        theme_layout.addWidget(self.article_display_checkbox)

        self.name_display_checkbox = ImageCheckBox("Имя")
        self.name_display_checkbox.setChecked(self.name_column_visible)
        self.name_display_checkbox.stateChanged.connect(self.toggle_name_column)
        theme_layout.addWidget(self.name_display_checkbox)

        self.total_qty_display_checkbox = ImageCheckBox("Всего")
        self.total_qty_display_checkbox.setChecked(True)  # Значение по умолчанию
        self.total_qty_display_checkbox.stateChanged.connect(self.toggle_total_qty_column)
        theme_layout.addWidget(self.total_qty_display_checkbox)

        self.stock_display_checkbox = ImageCheckBox("На складе")
        self.stock_display_checkbox.setChecked(True)  # Значение по умолчанию, будет применено позже из настроек
        self.stock_display_checkbox.stateChanged.connect(self.toggle_stock_column)
        theme_layout.addWidget(self.stock_display_checkbox)

        line = QFrame()
        line.setFrameShape(QFrame.Shape.VLine)
        line.setFrameShadow(QFrame.Shadow.Sunken)
        theme_layout.addWidget(line)

        self.hide_completed_checkbox = ImageCheckBox("Скрывать собранное")
        self.hide_completed_checkbox.setChecked(False)  # Значение по умолчанию, будет применено позже из настроек
        self.hide_completed_checkbox.stateChanged.connect(self.toggle_hide_completed_rows)
        theme_layout.addWidget(self.hide_completed_checkbox)

        # Применяем стили к чекбоксам с изображениями
        self.apply_toggle_styles()

        top_layout.addLayout(buttons_layout)
        top_layout.addLayout(theme_layout)

        return top_layout

    def _apply_visibility_settings_to_checkboxes(self):
        """Применить загруженные настройки видимости к чекбоксам и столбцам"""
        self.logger.info(f"_apply_visibility_settings_to_checkboxes вызвана")
        
        if not hasattr(self, 'current_user') or not self.current_user:
            self.logger.warning("current_user не установлен при применении настроек видимости")
            return

        # Используем сохранённые переменные экземпляра
        article_visible = getattr(self, 'article_column_visible', True)
        name_visible = getattr(self, 'name_column_visible', False)
        total_qty_visible = getattr(self, 'total_qty_column_visible', True)
        stock_visible = getattr(self, 'stock_column_visible', True)
        hide_completed = getattr(self, 'hide_completed_items_setting', False)

        self.logger.info(f"Применение настроек видимости: article={article_visible}, name={name_visible}, total_qty={total_qty_visible}, stock={stock_visible}, hide_completed={hide_completed}")
        self.logger.info(f"Состояние чекбоксов ДО применения: stock_display_checkbox.isChecked()={self.stock_display_checkbox.isChecked() if hasattr(self, 'stock_display_checkbox') else 'N/A'}, hide_completed_checkbox.isChecked()={self.hide_completed_checkbox.isChecked() if hasattr(self, 'hide_completed_checkbox') else 'N/A'}")

        # Обновляем состояние чекбоксов без вызова сигналов
        if hasattr(self, 'article_display_checkbox'):
            self.article_display_checkbox.blockSignals(True)
            self.article_display_checkbox.setChecked(article_visible)
            self.article_display_checkbox.blockSignals(False)
            # Принудительно обновляем иконку
            if hasattr(self.article_display_checkbox, '_update_icon'):
                self.article_display_checkbox._update_icon()

        if hasattr(self, 'name_display_checkbox'):
            self.name_display_checkbox.blockSignals(True)
            self.name_display_checkbox.setChecked(name_visible)
            self.name_display_checkbox.blockSignals(False)
            # Принудительно обновляем иконку
            if hasattr(self.name_display_checkbox, '_update_icon'):
                self.name_display_checkbox._update_icon()

        if hasattr(self, 'total_qty_display_checkbox'):
            self.total_qty_display_checkbox.blockSignals(True)
            self.total_qty_display_checkbox.setChecked(total_qty_visible)
            self.total_qty_display_checkbox.blockSignals(False)
            # Принудительно обновляем иконку
            if hasattr(self.total_qty_display_checkbox, '_update_icon'):
                self.total_qty_display_checkbox._update_icon()

        if hasattr(self, 'stock_display_checkbox'):
            self.stock_display_checkbox.blockSignals(True)
            self.stock_display_checkbox.setChecked(stock_visible)
            self.stock_display_checkbox.blockSignals(False)
            self.stock_display_checkbox.update()  # Принудительная перерисовка
            self.logger.debug(f"Установлено stock_display_checkbox.setChecked({stock_visible})")
            # Принудительно обновляем иконку
            if hasattr(self.stock_display_checkbox, '_update_icon'):
                self.stock_display_checkbox._update_icon()

        if hasattr(self, 'hide_completed_checkbox'):
            self.hide_completed_checkbox.blockSignals(True)
            self.hide_completed_checkbox.setChecked(hide_completed)
            self.hide_completed_checkbox.blockSignals(False)
            self.hide_completed_checkbox.update()  # Принудительная перерисовка
            self.logger.debug(f"Установлено hide_completed_checkbox.setChecked({hide_completed})")
            # Принудительно обновляем иконку
            if hasattr(self.hide_completed_checkbox, '_update_icon'):
                self.hide_completed_checkbox._update_icon()

        self.logger.info(f"Состояние чекбоксов ПОСЛЕ применения: stock_display_checkbox.isChecked()={self.stock_display_checkbox.isChecked() if hasattr(self, 'stock_display_checkbox') else 'N/A'}, hide_completed_checkbox.isChecked()={self.hide_completed_checkbox.isChecked() if hasattr(self, 'hide_completed_checkbox') else 'N/A'}")

        # Применяем видимость к столбцам таблицы
        if hasattr(self, 'shipment_table') and self.shipment_table:
            self.shipment_table.setColumnHidden(ColumnIndex.SKU, not article_visible)
            self.shipment_table.setColumnHidden(ColumnIndex.NAME, not name_visible)
            self.shipment_table.setColumnHidden(ColumnIndex.TOTAL_QTY, not total_qty_visible)
            self.shipment_table.setColumnHidden(ColumnIndex.STOCK_QTY, not stock_visible)

    def _init_main_splitter(self):
        """Создание основного горизонтального сплиттера"""
        return QSplitter(Qt.Orientation.Horizontal)

    def _init_left_panel(self, main_splitter):
        """Создание левой панели со списком поставок (дерево)"""
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.addWidget(QLabel("Поставки и коробки:"))

        self.shipments_tree_widget = QTreeWidget()
        self.shipments_tree_widget.setHeaderHidden(True)
        self.shipments_tree_widget.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.shipments_tree_widget.customContextMenuRequested.connect(self.show_shipment_context_menu)
        self.shipments_tree_widget.itemClicked.connect(self.on_shipment_clicked)
        self.shipments_tree_widget.itemDoubleClicked.connect(self.on_shipment_double_clicked)
        self.shipments_tree_widget.itemExpanded.connect(self.on_shipment_expanded)
        self.shipments_tree_widget.itemCollapsed.connect(self.on_shipment_collapsed)

        self.shipments_tree_widget.setMinimumWidth(150)
        self.shipments_tree_widget.setMaximumWidth(500)
        self.shipments_tree_widget.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        self.shipments_tree_widget.setIndentation(8)  # Уменьшили отступ иерархии
        self.shipments_tree_widget.setStyleSheet("""
            QTreeWidget {
                outline: 0;
                background-color: transparent;
            }
            QTreeWidget::item {
                border: 1px solid transparent;
                border-radius: 4px;
                margin: 2px;
                padding: 2px;
            }
            QTreeWidget::item:hover {
                background-color: rgba(0, 0, 0, 10);
            }
            QTreeWidget::item:selected {
                background-color: rgba(76, 175, 80, 40);
                border: 1px solid rgba(76, 175, 80, 120);
            }
        """)

        left_layout.addWidget(self.shipments_tree_widget)
        main_splitter.addWidget(left_widget)

    def _init_center_panel(self, main_splitter):
        """Создание центральной панели с таблицей текущей коробки и полем сканирования"""
        center_widget = QWidget()
        center_layout = QVBoxLayout(center_widget)
        center_layout.setContentsMargins(0, 0, 0, 0)

        self.current_box_label = QLabel("Коробка не выбрана")
        center_layout.addWidget(self.current_box_label)

        self.current_box_table = CustomTableWidget(0, 3)
        self.current_box_table.setAlternatingRowColors(True)
        self.current_box_table.setHorizontalHeaderLabels(["Штрихкод", "Артикул", "Кол-во"])
        self.current_box_table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked)
        # Стили будут применены через тему в apply_settings()

        for col in range(2):
            for row in range(self.current_box_table.rowCount()):
                item = self.current_box_table.item(row, col)
                if item:
                    item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)

        self.current_box_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.current_box_table.customContextMenuRequested.connect(self.show_box_table_context_menu)
        self.current_box_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self.current_box_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.current_box_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        self.current_box_table.itemDoubleClicked.connect(self.on_box_table_item_double_clicked)
        self.current_box_table.horizontalHeader().sectionResized.connect(self.on_table_columns_resized)

        center_layout.addWidget(self.current_box_table)

        # Создаем поле ввода штрихкода с иконкой слева
        scan_widget = QWidget()
        scan_layout = QHBoxLayout(scan_widget)
        scan_layout.setContentsMargins(0, 0, 0, 0)
        scan_layout.setSpacing(6)

        # Иконка штрихкода (используем кэш)
        barcode_icon = QLabel()
        from image_cache import get_cached_pixmap
        from app_constants import ICON_SIZE_MEDIUM
        barcode_pixmap = get_cached_pixmap(
            config.get_resource_path(Path("Res") / "basrc.png"),
            (ICON_SIZE_MEDIUM, ICON_SIZE_MEDIUM)
        )
        barcode_icon.setPixmap(barcode_pixmap)
        barcode_icon.setFixedSize(24, 24)
        barcode_icon.setAlignment(Qt.AlignmentFlag.AlignCenter)
        scan_layout.addWidget(barcode_icon)

        # Поле ввода штрихкода
        self.scan_input = QLineEdit()
        self.scan_input.setPlaceholderText("Отсканируйте штрихкод и нажмите Enter")
        self.scan_input.returnPressed.connect(self.handle_scan)
        self.scan_input.setStyleSheet(f"""
            QLineEdit {{
                font-size: 16px;
                padding: 8px 12px;
                border: 2px solid #cccccc;
                border-radius: 8px;
                background-color: white;
            }}
            QLineEdit:focus {{
                border: 2px solid #007bff;
                background-color: #f8f8ff;
            }}
        """)
        scan_layout.addWidget(self.scan_input, 1)  # 1 = растягивать

        center_layout.addWidget(scan_widget)

        main_splitter.addWidget(center_widget)

    def _init_right_panel(self, main_splitter):
        """Создание правой панели с таблицей списка поставки"""
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(0, 0, 0, 0)

        self.shipment_table_label = QLabel("Спостав поставки:")
        right_layout.addWidget(self.shipment_table_label)

        # Таблица поставки: 6 столбцов
        # 0=Штрихкод, 1=Артикул, 2=Имя(скрыт), 3=Всего, 4=Осталось, 5=Склад, 6=Кнопка "+"
        self.shipment_table = CustomTableWidget(0, 7)
        self.shipment_table.setAlternatingRowColors(True)
        self.shipment_table.setHorizontalHeaderLabels(["Штрихкод", "Артикул", "Имя", "Всего", "Осталось", "Склад", ""])
        self.shipment_table.verticalHeader().setVisible(False)
        self.shipment_table.setSortingEnabled(True)
        self.shipment_table.cellChanged.connect(self.on_shipment_cell_changed)
        self.shipment_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)

        header = self.shipment_table.horizontalHeader()
        for col in range(self.shipment_table.columnCount()):
            header.setMinimumSectionSize(30)

        self.shipment_table.setColumnHidden(ColumnIndex.NAME, True)
        if self.shipment_table.columnCount() > ColumnIndex.REMAINING_QTY:
            self.shipment_table.setItemDelegateForColumn(ColumnIndex.REMAINING_QTY, QuantityEditDelegate())

        self.shipment_table.itemDoubleClicked.connect(self.on_shipment_table_item_double_clicked)
        self.shipment_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.shipment_table.customContextMenuRequested.connect(self.show_shipment_table_context_menu)
        self.shipment_table.horizontalHeader().sectionResized.connect(self.on_table_columns_resized)
        self.shipment_table.horizontalHeader().sortIndicatorChanged.connect(self.on_shipment_table_sort_changed)

        # Добавляем кнопку "+ Все" в заголовок столбца "Действие" (столбец 6)
        self._init_add_all_button_in_header()

        right_layout.addWidget(self.shipment_table)
        right_widget.setMinimumWidth(400)
        main_splitter.addWidget(right_widget)
        return right_layout

    def _init_removed_items_section(self, right_layout):
        """Создание секции для отображения удалённых из поставки items"""
        self.removed_items_label = QLabel("Убрать из поставки:")
        self.removed_items_label.setStyleSheet("font-weight: bold; color: #cc0000; margin-top: 10px;")
        self.removed_items_label.setVisible(False)
        right_layout.addWidget(self.removed_items_label)

        self.removed_items_table = CustomTableWidget(0, 3)
        self.removed_items_table.setHorizontalHeaderLabels(["Штрихкод", "Артикул", "Убрать из коробки"])
        self.removed_items_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        self.removed_items_table.verticalHeader().setVisible(False)
        self.removed_items_table.setVisible(False)
        self.removed_items_table.horizontalHeader().sectionResized.connect(self.on_table_columns_resized)
        right_layout.addWidget(self.removed_items_table)

    def _init_add_all_button_in_header(self):
        """Добавить кнопку '+ Все' в заголовок столбца 'Действие'"""
        from PyQt6.QtWidgets import QPushButton, QWidget
        from PyQt6.QtCore import Qt
        
        # Создаем контейнер для кнопки в заголовке
        header = self.shipment_table.horizontalHeader()
        
        # Создаем кнопку "+ Все"
        self.add_all_header_btn = QPushButton("+ Все")
        self.add_all_header_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                    stop: 0 #ff9800,
                    stop: 0.5 #ff9800,
                    stop: 1 #f57c00);
                color: white;
                border: 1px solid #e65100;
                padding: 1px 4px;
                border-radius: 2px;
                font-weight: bold;
                font-size: 8px;
                min-height: 14px;
                max-height: 14px;
            }
            QPushButton:hover {
                background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                    stop: 0 #ffb74d,
                    stop: 0.5 #ffa726,
                    stop: 1 #ff9800);
            }
            QPushButton:pressed {
                background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                    stop: 0 #f57c00,
                    stop: 0.5 #e65100,
                    stop: 1 #d84315);
            }
            QPushButton:disabled {
                background: #bdbdbd;
                color: #757575;
                border: 1px solid #9e9e9e;
            }
        """)
        self.add_all_header_btn.setToolTip("Добавить весь остаток всех товаров в текущую коробку")
        self.add_all_header_btn.clicked.connect(self.on_add_all_button_clicked)
        self.add_all_header_btn.setParent(header)
        
        # Устанавливаем отступы для заголовка, чтобы освободить место для кнопки
        # Кнопка будет позиционирована вручную
        self.add_all_header_btn.show()

        # Обновляем позицию кнопки при изменении размера заголовка
        header.sectionResized.connect(self._update_add_all_button_position)
        self.shipment_table.horizontalHeader().sectionMoved.connect(self._update_add_all_button_position)

        # Первоначальное позиционирование и видимость кнопки
        self._update_add_all_button_position()
        self._update_add_all_button_visibility()
    
    def _update_add_all_button_position(self):
        """Обновить позицию кнопки '+ Все' в заголовке"""
        if not hasattr(self, 'add_all_header_btn'):
            return

        header = self.shipment_table.horizontalHeader()
        column_index = 6  # Столбец "Действие"

        # Получаем координаты заголовка столбца
        left = header.sectionPosition(column_index)
        right = left + header.sectionSize(column_index)
        top = 0
        height = header.height()

        # Вычисляем позицию кнопки (центрируем в заголовке)
        btn_width = self.add_all_header_btn.sizeHint().width()
        btn_height = self.add_all_header_btn.sizeHint().height()

        x = left + (right - left - btn_width) // 2
        y = (height - btn_height) // 2

        self.add_all_header_btn.setGeometry(x, y, btn_width, btn_height)

    def _update_add_all_button_visibility(self):
        """Обновить видимость кнопки '+ Все' в заголовке
        
        Кнопка отображается только когда:
        1. Выбрана одна из коробок (current_box_index >= 0)
        2. Есть нераспределённый товар (remaining_qty > 0 хотя бы у одного товара)
        """
        if not hasattr(self, 'add_all_header_btn'):
            return

        # Проверяем, выбрана ли коробка
        if not self.current_shipment or self.current_shipment.current_box_index < 0:
            self.add_all_header_btn.hide()
            return

        # Проверяем, есть ли нераспределённый товар
        has_unallocated = False
        for barcode, shipment_item in self.current_shipment.shipment_items.items():
            # Пропускаем удалённые товары
            if barcode in self.current_shipment.removed_items:
                continue
            if shipment_item.remaining_qty > 0:
                has_unallocated = True
                break

        if has_unallocated:
            self.add_all_header_btn.show()
        else:
            self.add_all_header_btn.hide()

    def _init_splitter_final_config(self, main_splitter, main_layout):
        """Финальная настройка сплиттера и добавление в main layout"""
        main_splitter.splitterMoved.connect(self.on_splitter_moved)
        main_splitter.setSizes([200, 300, 800])
        main_layout.addWidget(main_splitter, 1)

    def _init_timers_and_focus(self):
        """Инициализация таймеров, установка фокуса и фильтров событий"""
        self.cleanup_sessions_timer = QTimer()
        self.cleanup_sessions_timer.start(6000)
        self.scan_input.setFocus()
        self.update_scan_input_style()
        self.installEventFilter(self)
        self.current_box_table.installEventFilter(self)
        self.shipment_table.installEventFilter(self)
        self.removed_items_table.installEventFilter(self)
        self.shipments_tree_widget.installEventFilter(self)

    def init_ui(self):
        """Инициализация пользовательского интерфейса"""
        self._init_menu_bar()
        main_layout = self._init_main_layout()
        top_layout, buttons_layout = self._init_top_buttons_panel()
        top_layout = self._init_display_controls(top_layout, buttons_layout)
        main_layout.addLayout(top_layout)

        self.main_splitter = self._init_main_splitter()
        self._init_left_panel(self.main_splitter)
        self._init_center_panel(self.main_splitter)
        right_layout = self._init_right_panel(self.main_splitter)
        self._init_removed_items_section(right_layout)
        self._init_splitter_final_config(self.main_splitter, main_layout)
        self._init_timers_and_focus()

    def setup_shortcuts(self):
        # Шорткаты теперь обрабатываются через MVC контроллер в mvc_controller.py
        # QShortcut(QKeySequence("F1"), self).activated.connect(self.shipment_controller.handle_shipment_operations)
        # QShortcut(QKeySequence("F2"), self).activated.connect(self.shipment_controller.handle_group_shipment_operations)
        # QShortcut(QKeySequence("F3"), self).activated.connect(self.shipment_controller.handle_new_box)
        # QShortcut(QKeySequence("F5"), self).activated.connect(self.ui_controller.handle_refresh)
        # QShortcut(QKeySequence("F11"), self).activated.connect(self.ui_controller.handle_theme_shortcut)
        # QShortcut(QKeySequence("Ctrl+S"), self).activated.connect(self.ui_controller.handle_save_session)
        # QShortcut(QKeySequence("Ctrl+P"), self).activated.connect(self.ui_controller.handle_print_label)
        # QShortcut(QKeySequence("Ctrl+H"), self).activated.connect(self.ui_controller.handle_shipment_check) # Шорткат для "Проверка поставки"
        
        # Добавляем шорткат для пробела, чтобы перевести фокус на поле ввода штрихкода
        from PyQt6.QtGui import QKeySequence, QShortcut
        space_shortcut = QShortcut(QKeySequence("Space"), self)
        space_shortcut.activated.connect(self.focus_scan_input_field)
        
        # Теперь все шорткаты управляются через MVC контроллер, поставкляем этот метод для обратной совместимости
        pass

    def toggle_theme_shortcut(self):
        """Toggle the theme between light and dark"""
        if self.current_theme == "Светлая":
            self.current_theme = "Тёмная"
        elif self.current_theme == "Тёмная":
            self.current_theme = "macOS"
        else:
            self.current_theme = "Светлая"
        self.apply_settings()
        self.save_user_settings()

    def on_box_table_item_double_clicked(self, item):
        """
        Handle double click on box table items - open dialog for quantity editing
        """
        # Only handle editing for the quantity column (index 2)
        if item.column() == 2:
            self.open_quantity_edit_dialog(item.row())

    def open_quantity_edit_dialog(self, row):
        """
        Open dialog to edit quantity for the specified row
        """
        if not self.current_shipment or self.current_shipment.current_box_index < 0:
            return
            
        # Get the current box
        current_box = self.current_shipment.boxes[self.current_shipment.current_box_index]
        
        # Check if row is valid
        if row >= len(current_box.items):
            return
            
        # Get the keys of items in the current box
        keys = list(current_box.items.keys())
        if row >= len(keys):
            return
            
        # Get the barcode for this row
        barcode = keys[row]
        
        # Get the current quantity
        current_qty = current_box.items[barcode]
        
        # Create and show the dialog
        dialog = QuantityEditDialog(current_qty, self)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            new_qty = dialog.get_quantity()
            
            # Validate the new quantity
            shipment_item = self.current_shipment.shipment_items.get(barcode)
            
            if barcode in self.current_shipment.removed_items:
                if new_qty > current_box.items.get(barcode, 0):
                    from PyQt6.QtWidgets import QMessageBox
                    QMessageBox.warning(
                        self, "Ошибка",
                        f"Товар «{barcode}» удален из поставки! Можно только уменьшать количество."
                    )
                    return
            
            if shipment_item:
                old_qty = current_box.items.get(barcode, 0)
                diff = new_qty - old_qty
                if shipment_item.allocated_qty + diff > shipment_item.total_qty:
                    from PyQt6.QtWidgets import QMessageBox
                    QMessageBox.warning(
                        self, "Ошибка",
                        f"Нельзя распределить больше {shipment_item.total_qty} единиц товара {barcode}!"
                    )
                    return
                shipment_item.allocated_qty += diff
            
            # Update the quantity in the box
            current_box.set_item_qty(barcode, new_qty)
            
            # Update removed_items if this item is in the removed list
            if barcode in self.current_shipment.removed_items:
                removed_qty = self.current_shipment.removed_items[barcode]['allocated_qty']
                new_removed_qty = removed_qty - diff  # When we reduce quantity in box, removed qty increases by the same amount
                
                if new_removed_qty <= 0:
                    del self.current_shipment.removed_items[barcode]
                else:
                    self.current_shipment.removed_items[barcode]['allocated_qty'] = new_removed_qty
            elif new_qty == 0 and barcode in self.current_shipment.removed_items:
                del self.current_shipment.removed_items[barcode]
            
            # Check if this item should be added to removed_items due to total quantity exceeded
            # This can happen if the total quantity in shipment was reduced but this box still has more than allowed
            if (shipment_item and
                shipment_item.allocated_qty > shipment_item.total_qty and
                barcode not in self.current_shipment.removed_items):
                excess_qty = shipment_item.allocated_qty - shipment_item.total_qty
                self.current_shipment.removed_items[barcode] = {
                    'sku': shipment_item.sku,
                    'allocated_qty': excess_qty
                }
            
            # Also check if we need to remove item from removed_items if it's no longer in excess
            if (barcode in self.current_shipment.removed_items and
                shipment_item and
                shipment_item.allocated_qty <= shipment_item.total_qty):
                del self.current_shipment.removed_items[barcode]
            
            # Сбрасываем кэши в модели поставки при изменении распределения товаров
            self.current_shipment.invalidate_caches()

            # Update UI
            self.ui_updater.update_current_components()
            self.ui_updater.update_shipments_tree()
            
            # Show status message
            # Статус скрыт, сообщения не отображаются
            # self.statusBar().showMessage(f"Количество товара {barcode} изменено на {new_qty}", 3000)
            
            # Schedule a delayed save to the database
            if hasattr(self.shipment_manager, 'schedule_save'):
                self.shipment_manager.schedule_save()
    
    def on_shipment_table_item_double_clicked(self, item):
        """
        Handle double click on shipment table items to adjust column width during editing
        """
        # Only adjust width for the total column (index 2)
        if item.column() == 2:
            # Temporarily increase column width to accommodate editing
            # Calculate text width to ensure it fits properly
            text = item.text()
            font_metrics = self.shipment_table.fontMetrics()
            text_width = font_metrics.horizontalAdvance(text + "00")  # Add some padding
            min_width = max(self.shipment_table.columnWidth(item.column()), text_width, 10)
            self.shipment_table.setColumnWidth(item.column(), min_width)

    def on_shipment_table_sort_changed(self, logical_index, order):
        """Обработкатчик изменения сортировки таблицы поставки - обновляет видимость строк"""
        if hasattr(self, 'ui_updater') and self.ui_updater:
            self.ui_updater.update_shipment_table_rows_visibility()

    def on_add_all_button_clicked(self):
        """Обработчик нажатия кнопки '+ Все' в заголовке столбца"""
        if hasattr(self, 'shipment_manager') and self.shipment_manager:
            self.shipment_manager.add_all_remaining_for_all_items_to_box()
            # Обновляем видимость кнопки после добавления товаров
            self._update_add_all_button_visibility()

    def update_scan_input_style(self):
        try:
            theme = themes.THEMES.get(self.current_theme, themes.THEMES["Светлая"])
            border_color = theme["input_border"].name()
            focus_color = theme["accent_primary"].name()
            bg_color = theme["input_bg"].name()

            if hasattr(self, 'scan_input') and self.scan_input:
                self.scan_input.setStyleSheet(f"""
                    QLineEdit {{
                        font-size: 16px;
                        padding: 8px 12px;
                        border: 2px solid {border_color};
                        border-radius: 8px;
                        background-color: {bg_color};
                        color: {theme["input_text"].name()};
                    }}
                    QLineEdit:focus {{
                        border: 2px solid {focus_color};
                        background-color: {"#f8f8ff" if self.current_theme == "Светлая" else "#444454"};
                    }}
                """)
        except Exception as e:
            self.logger.error(f"Ошибка при обновлении стиля поля сканирования: {e}", exc_info=True)

    def on_table_columns_resized(self):
        """Обработчик изменения размера столбцов таблиц"""
        # Сохраняем ширину столбцов только если инициализация завершена
        if getattr(self, 'initialization_complete', False):
            self.save_columns_width()
            self.schedule_save_settings()

    def on_splitter_moved(self):
        """Обработкатчик изменения размера сплиттера"""
        self.save_window_state()
        self.schedule_save_settings()

    def schedule_save_settings(self):
        """Запланировать сохранение настроек с задержкой"""
        self.save_timer.start(100)

    def save_columns_width(self):
        """Сохранить ширину столбцов таблиц"""
        # Сохраняем ширину столбцов для таблицы поставки
        if hasattr(self, 'shipment_table') and self.shipment_table:
            shipment_widths = []
            for col in range(self.shipment_table.columnCount()):
                shipment_widths.append(str(self.shipment_table.columnWidth(col)))
            self.shipment_columns_width = ",".join(shipment_widths)
        
        # Сохраняем ширину столбцов для таблицы коробки
        if hasattr(self, 'current_box_table') and self.current_box_table:
            box_widths = []
            for col in range(self.current_box_table.columnCount()):
                box_widths.append(str(self.current_box_table.columnWidth(col)))
            self.box_columns_width = ",".join(box_widths)
        
        # Сохраняем ширину столбцов для таблицы убранных из поставки товаров
        # (не требуется сохранять отдельно, т.к. используется общее поле, но если потребуется, можно добавить отдельное поле)

    def toggle_name_column(self, state):
        """Управление отображением столбца "Имя"""
        try:
            is_visible = (state == Qt.CheckState.Checked.value)
            self.name_column_visible = is_visible
            self.shipment_table.setColumnHidden(ColumnIndex.NAME, not is_visible)

            if is_visible and self.shipment_table.columnWidth(ColumnIndex.NAME) == 0:
                self.shipment_table.setColumnWidth(ColumnIndex.NAME, 100)

            if getattr(self, 'initialization_complete', False) and not getattr(self, '_loading_settings', False):
                self.save_user_settings()
        except Exception as e:
            self.logger.error(f"Ошибка при переключении столбца 'Имя': {e}", exc_info=True)

    def toggle_total_qty_column(self, state):
        """Управление отображением столбца "Всего"""
        try:
            is_visible = (state == Qt.CheckState.Checked.value)
            self.total_qty_column_visible = is_visible
            self.shipment_table.setColumnHidden(ColumnIndex.TOTAL_QTY, not is_visible)

            if is_visible and self.shipment_table.columnWidth(ColumnIndex.TOTAL_QTY) == 0:
                self.shipment_table.setColumnWidth(ColumnIndex.TOTAL_QTY, 100)

            if getattr(self, 'initialization_complete', False) and not getattr(self, '_loading_settings', False):
                self.save_user_settings()
        except Exception as e:
            self.logger.error(f"Ошибка при переключении столбца 'Всего': {e}", exc_info=True)
            
    def toggle_article_column(self, state):
        """Управление отображением столбца "Артикул"""
        try:
            is_visible = (state == Qt.CheckState.Checked.value)
            self.article_column_visible = is_visible
            self.shipment_table.setColumnHidden(ColumnIndex.SKU, not is_visible)

            if is_visible and self.shipment_table.columnWidth(ColumnIndex.SKU) == 0:
                self.shipment_table.setColumnWidth(ColumnIndex.SKU, 100)

            init_complete = getattr(self, 'initialization_complete', False)
            loading = getattr(self, '_loading_settings', False)
            self.logger.debug(f"toggle_article_column: init={init_complete}, loading={loading}, save={init_complete and not loading}, article={self.article_column_visible}")
            if init_complete and not loading:
                self.save_user_settings()
        except Exception as e:
            self.logger.error(f"Ошибка при переключении столбца 'Артикул': {e}", exc_info=True)
            
    def toggle_stock_column(self, state):
        """Управление отображением столбца "На складе"""
        try:
            is_visible = (state == Qt.CheckState.Checked.value)
            self.stock_column_visible = is_visible
            self.shipment_table.setColumnHidden(ColumnIndex.STOCK_QTY, not is_visible)

            if is_visible and self.shipment_table.columnWidth(ColumnIndex.STOCK_QTY) == 0:
                self.shipment_table.setColumnWidth(ColumnIndex.STOCK_QTY, 100)

            if getattr(self, 'initialization_complete', False) and not getattr(self, '_loading_settings', False):
                self.save_user_settings()
        except Exception as e:
            self.logger.error(f"Ошибка при переключении столбца 'На складе': {e}", exc_info=True)

    def toggle_hide_completed_rows(self, state):
        """Управление скрытием полностью собранных строк (где 'Осталось' = 0) в таблице поставки"""
        try:
            hide_completed = bool(state)

            if hasattr(self, 'current_shipment') and self.current_shipment:
                self.current_shipment.hide_completed_items = hide_completed

            self.hide_completed_items_setting = hide_completed

            if getattr(self, 'initialization_complete', False) and not getattr(self, '_loading_settings', False):
                self.save_user_settings()

            if hasattr(self, 'ui_updater') and self.ui_updater:
                self.ui_updater.update_shipment_table_rows_visibility()
        except Exception as e:
            self.logger.error(f"Ошибка при переключении скрытия собранных строк: {e}", exc_info=True)

    def update_user_display(self):
        """Обновление отображения пользователя в меню"""
        if self.current_user:
            self.user_action.setText(f"👤 {self.current_user}")
        else:
            self.user_action.setText("👤 Выбор пользователя")

    def manage_users(self):
        try:
            dialog = UserManagerDialog(self.current_user, self)
            if dialog.exec() == QDialog.DialogCode.Accepted:
                selected_user = dialog.get_selected_user()
                if selected_user and selected_user != self.current_user:
                    self.save_user_settings()
                    self.current_user = selected_user
                    utils.save_local_user(self.current_user)
                    self.load_user_settings()
                    self.apply_settings()
                    # Загружаем состояние окна для нового пользователя
                    self.load_window_state()
                    self.update_user_display()
                    # Статус скрыт, сообщения не отображаются
                    # self.statusBar().showMessage(f"Пользователь изменен на: {self.current_user}", 300)
        except Exception as e:
            self.logger.error(f"Ошибка при управлении пользователями: {e}", exc_info=True)
            QMessageBox.critical(self, "Ошибка", f"Ошибка при управлении пользователями:\n{e}")

    def load_window_state(self):
        """Загрузка размеров окна и сплиттера из настроек пользователя"""
        if self.window_width and self.window_height:
            self.resize(self.window_width, self.window_height)
        
        # Отложенная загрузка размеров сплиттера, чтобы убедиться, что UI полностью загружен
        if hasattr(self, 'main_splitter') and self.main_splitter_sizes:
            try:
                sizes = [int(x) for x in self.main_splitter_sizes.split(",")]
                if len(sizes) == 3:
                    if sizes[0] < 100:
                        sizes[0] = 100
                    if sizes[2] < 400:
                        sizes[2] = 400  # Исправлено: было 800, что может быть слишком большешим
                    # Используем QTimer для отложенной установки размеров сплиттера
                    # Увеличиваем задержку, чтобы гарантировать полную загрузку UI
                    QTimer.singleShot(300, lambda: self.main_splitter.setSizes(sizes))
                    # Принудительно обновляем геометрию сплиттера после установки размеров
                    QTimer.singleShot(350, lambda: self.main_splitter.updateGeometry())
                    # Дополнительно обновляем все компоненты сплиттера
                    QTimer.singleShot(400, lambda: self.update_splitter_widgets_geometry())
            except (ValueError, AttributeError):
                # Если не удалось восстановить сохраненные размеры, устанавливаем стандартные
                QTimer.singleShot(300, lambda: self.main_splitter.setSizes([200, 300, 800]))
                # Принудительно обновляем геометрию сплиттера
                QTimer.singleShot(350, lambda: self.main_splitter.updateGeometry())
                # Дополнительно обновляем все компоненты сплиттера
                QTimer.singleShot(400, lambda: self.update_splitter_widgets_geometry())
        
        # Дополнительно обновляем таблицы для гарантии отображения
        if hasattr(self, 'shipment_table') and self.shipment_table:
            self.shipment_table.updateGeometry()
        if hasattr(self, 'current_box_table') and self.current_box_table:
            self.current_box_table.updateGeometry()

        # Обновляем UI для гарантии отображения всех данных
        QTimer.singleShot(400, self.update_ui)

        self.logger.info("Завершение загрузки состояния окна")
        
    def update_splitter_widgets_geometry(self):
        """Обновить геометрию всех виджетов в сплиттере"""
        if hasattr(self, 'main_splitter'):
            for i in range(self.main_splitter.count()):
                widget = self.main_splitter.widget(i)
                if widget:
                    widget.updateGeometry()
                    widget.update()

    def save_window_state(self):
        """Сохранение размеров окна и сплиттера в настройки пользователя"""
        self.window_width = self.width()
        self.window_height = self.height()
        
        if hasattr(self, 'main_splitter'):
            # Сохраняем текущующие размеры сплиттера
            current_sizes = self.main_splitter.sizes()
            self.main_splitter_sizes = ",".join(str(s) for s in current_sizes)

    def closeEvent(self, event):
        try:
            self.save_timer.stop()
            # Принудительно сохраняем ширину столбцов перед сохранением настроек
            self.save_columns_width()
            self.save_window_state()
            # Сохраняем ВСЕ поставки, а не только текущую (пользователь мог редактировать несколько)
            self.save_all_shipments()
            # Сохраняем настройки пользователя
            self.save_user_settings()
            self.logger.info(f"Настройки пользователя {self.current_user} сохранены при закрытии")
            # Cancel all active async operations before closing
            if self.async_manager is not None:
                self.async_manager.cancel_all_operations()
            else:
                self.logger.debug("async_manager еще не инициализирован, пропускаем отмену операций")
            # Cleanup Moysklad sync thread
            if hasattr(self, 'improved_sync_handler') and self.improved_sync_handler:
                self.improved_sync_handler.cleanup()
            event.accept()
        except Exception as e:
            self.logger.error(f"Ошибка при закрытии приложения: {e}", exc_info=True)
            event.accept()  # В любом случае разрешаем зкакрытие приложения

    def initialize_user(self):
        users = database.get_all_users()
        if not users:
            try:
                database.set_user_settings("Default", config.DEFAULT_FONT_SIZE,
                                         config.DEFAULT_LABEL_FONT_SIZE, config.DEFAULT_THEME,
                                         "ok.wav", "error.wav", "", "", "", 1300, 800, "", "", "", "",
                                         "", "[]", True, True, False, True, "", True, "")  # Добавляем значения по умолчанию для всех параметров
                users = database.get_all_users()
            except Exception as e:
                import logging
                logger = logging.getLogger(__name__)
                logger.error(f"Ошибка создания пользователя по умолчанию: {e}")
                self.current_user = "Default"
                return
        
        local_user = utils.load_local_user()
        # Check if users is a valid list and handle dict (PostgreSQL) results
        if users and isinstance(users, list):
            # Function to safely get username from dict or tuple
            def get_username(user_entry):
                if isinstance(user_entry, tuple):
                    # If it's a tuple (username, ...), return the first element
                    return user_entry[0] if len(user_entry) >= 1 else None
                elif hasattr(user_entry, '__getitem__'):  # It's a dict
                    try:
                        # For dict-like objects (RealDictCursor), access by key
                        return user_entry.get('username') if hasattr(user_entry, 'get') else (user_entry['username'] if 'username' in user_entry else None)
                    except (KeyError, TypeError, IndexError):
                        return None
                return None
            
            # Filter out any invalid entries
            valid_users = [u for u in users if get_username(u) is not None]
            if local_user and any(get_username(u) == local_user for u in valid_users):
                self.current_user = local_user
            else:
                username = get_username(valid_users[0]) if valid_users else None
                self.current_user = username if username else "Default"
        else:
            self.current_user = "Default"

        # Настройки загружаются в deferred_initialization() после создания UI

    def load_user_settings(self):
        if not self.current_user:
            self.logger.debug("load_user_settings: current_user не установлен")
            return
        settings = database.get_user_settings(self.current_user)
        if settings:
            self.logger.debug(
                f"Загрузка настроек пользователя {self.current_user}: "
                f"article_visible={settings.get('article_column_visible', True)}, "
                f"name_visible={settings.get('name_column_visible', False)}, "
                f"stock_visible={settings.get('stock_column_visible', True)}"
            )
            
            self.font_size = settings["font_size"]
            self.label_font_size = settings["label_font_size"]
            self.current_theme = settings["theme"]
            self.ok_sound = settings["ok_sound"]
            self.error_sound = settings["error_sound"]
            self.tone_sound = settings.get("tone_sound", False)
            self.sound_volume = settings.get("sound_volume", 100)
            
            # Применяем громкость звуков
            utils.set_sound_volume(self.sound_volume)
            
            self.shipment_columns_width = settings["shipment_columns_width"]
            self.box_columns_width = settings["box_columns_width"]
            self.main_splitter_sizes = settings["main_splitter_sizes"]
            self.window_width = settings["window_width"]
            self.window_height = settings["window_height"]
            # Загружаем сохранённые цвета кнопок
            import json
            button_colors_str = settings.get("button_colors", "{}")
            try:
                self.button_colors = json.loads(button_colors_str) if button_colors_str else {}
                self.logger.info(f"Загружены цвета кнопок")
            except Exception as e:
                self.logger.error(f"Ошибка загрузки цветов кнопок: {e}")
                self.button_colors = {}
            
            # Применяем загруженные цвета кнопок
            self.apply_button_colors()
            
            # Система блокировки ппоставккуавок удалена, поэтому не загружаем эту настройку
            # self.shipment_locking_enabled = settings["shipment_locking_enabled"]

            # Загружаем настройки видимости столбцов
            self.article_column_visible = settings.get("article_column_visible", True)
            self.name_column_visible = settings.get("name_column_visible", False)
            self.total_qty_column_visible = settings.get("total_qty_column_visible", True)
            self.stock_column_visible = settings.get("stock_column_visible", True)
            self.hide_completed_items_setting = settings.get("hide_completed_items", False)

            self.logger.info(f"Загруженные настройки видимости: article={self.article_column_visible}, name={self.name_column_visible}, total_qty={self.total_qty_column_visible}, stock={self.stock_column_visible}, hide_completed={self.hide_completed_items_setting}")
            self.logger.info(f"Загруженные настройки из БД: stock_column_visible={settings.get('stock_column_visible', True)}, hide_completed_items={settings.get('hide_completed_items', False)}")

            # Видимость столбцов и чекбоксов применяется в _apply_visibility_settings_to_checkboxes()

        # Обновляем состояние чекбокса "Скрывать собранное" в текущей поставке
        if hasattr(self, 'current_shipment') and self.current_shipment:
            self.current_shipment.hide_completed_items = self.hide_completed_items_setting

            # Применяем скрытие собранных строк, если есть текущующая поставка
            if hasattr(self, 'ui_updater') and self.ui_updater:
                self.ui_updater.update_shipment_table_rows_visibility()

            # Обновляем видимость кнопки МойСклад после загрузки настроек
            self.update_moysklad_button_visibility()

            # Применяем стили к чекбоксам после загрузки темы
            self.apply_toggle_styles()

    def restore_table_columns_width(self):
        """Восстановление ширины столбцов таблиц"""
        # Восстанавливаем сохраненные ширины столбцов, если они есть
        if hasattr(self, 'shipment_table') and self.shipment_table and self.shipment_columns_width:
            try:
                widths = self.shipment_columns_width.split(",")
                for i, width_str in enumerate(widths):
                    if i < self.shipment_table.columnCount():
                        width = int(width_str)
                        self.shipment_table.setColumnWidth(i, width)
                # Обновляем геометрию таблицы после восстановления ширины столбцов
                self.shipment_table.updateGeometry()
            except (ValueError, IndexError):
                # Если не удалось восстановить сохраненные ширины, используем авто-размер
                self.shipment_table.resizeColumnsToContents()
                self.shipment_table.updateGeometry()
        elif hasattr(self, 'shipment_table') and self.shipment_table:
            self.shipment_table.resizeColumnsToContents()
            self.shipment_table.updateGeometry()
        
        # Принудительно устанавливаем ширину столбца 6 с кнопками, т.к. resizeColumnsToContents не учитывает виджеты
        if hasattr(self, 'shipment_table') and self.shipment_table:
            self.shipment_table.setColumnWidth(ColumnIndex.ACTION, 100)
            self.shipment_table.setColumnHidden(ColumnIndex.ACTION, False)

        if hasattr(self, 'current_box_table') and self.current_box_table and self.box_columns_width:
            try:
                widths = self.box_columns_width.split(",")
                for i, width_str in enumerate(widths):
                    if i < self.current_box_table.columnCount():
                        width = int(width_str)
                        self.current_box_table.setColumnWidth(i, width)
                # Обновляем геометрию таблицы после восстановления ширины столбцов
                self.current_box_table.updateGeometry()
            except (ValueError, IndexError):
                # Если не удалось восстановить сохраненные ширины, испо��ьзуем авто-размер
                self.current_box_table.resizeColumnsToContents()
                self.current_box_table.updateGeometry()
        elif hasattr(self, 'current_box_table') and self.current_box_table:
            self.current_box_table.resizeColumnsToContents()
            self.current_box_table.updateGeometry()
        
        # Восстанавливаем режимы изменения размера заголовков
        if hasattr(self, 'shipment_table') and self.shipment_table:
            # Устанавливаем фиксированный режим для вертикального заголовка, чтобы сохранить нашу высоту строк
            try:
                from PyQt6.QtWidgets import QHeaderView
                self.shipment_table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Fixed)
            except:
                # Если QHeaderView недоступен, используем числовое значение (Fixed = 2)
                self.shipment_table.verticalHeader().setSectionResizeMode(2)
        if hasattr(self, 'current_box_table') and self.current_box_table:
            # Устанавливаем фиксированный режим для вертикального заголовка, чтобы сохранить нашу высоту строк
            try:
                from PyQt6.QtWidgets import QHeaderView
                self.current_box_table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Fixed)
            except:
                # Если QHeaderView недоступен, используем числовое значение (Fixed = 2)
                self.current_box_table.verticalHeader().setSectionResizeMode(2)
                
        # После восстановления сохраненных ширин большеше не вызываем resizeColumnsToContents
        self.logger.info("Завершение восстановления ширины столбцов таблиц")
            
    def save_user_settings(self):
        self.logger.info(f"save_user_settings вызван: initialization_complete={getattr(self, 'initialization_complete', False)}, _loading_settings={getattr(self, '_loading_settings', False)}")
        
        # Не сохраняем настройки во время инициализации
        if not getattr(self, 'initialization_complete', False):
            self.logger.warning("save_user_settings: initialization_complete=False, выход")
            return

        # Ткакже не сохраняем, если идёт загрузка настроек
        if getattr(self, '_loading_settings', False):
            self.logger.warning("save_user_settings: _loading_settings=True, выход")
            return
        
        # Не сохраняем, если чекбоксы ещё не созданы
        if not hasattr(self, 'article_display_checkbox'):
            return
        
        if self.current_user:
            self.save_columns_width()
            self.save_window_state()

            # Получаем текущующие настройки пользователя
            current_settings = database.get_user_settings(self.current_user)
            
            # Настройки МойСклад теперь глобальные (для всех пользователей)
            moysklad_token = database.get_moysklad_token()
            moysklad_stores = database.get_moysklad_stores() or '[]'
            moysklad_enabled = database.get_moysklad_enabled()
            colored_buttons = current_settings.get('colored_buttons', config.DEFAULT_COLORED_BUTTONS) if current_settings else config.DEFAULT_COLORED_BUTTONS

            # Используем сохранённые переменные экземпляра для видимости столбцов
            article_column_visible = getattr(self, 'article_column_visible', True)
            name_column_visible = getattr(self, 'name_column_visible', False)
            total_qty_column_visible = getattr(self, 'total_qty_column_visible', True)
            stock_column_visible = getattr(self, 'stock_column_visible', True)
            hide_completed_items = getattr(self, 'hide_completed_items_setting', False)

            self.logger.info(f"save_user_settings: Сохранение настроек видимости: article={article_column_visible}, name={name_column_visible}, total_qty={total_qty_column_visible}, stock={stock_column_visible}, hide_completed={hide_completed_items}")

            database.set_user_settings(
                self.current_user,
                self.font_size,
                self.label_font_size,
                self.current_theme,
                self.ok_sound,
                self.error_sound,
                self.shipment_columns_width,
                self.box_columns_width,
                self.main_splitter_sizes,
                self.window_width,
                self.window_height,
                "",  # button_primary_color
                "",  # button_success_color
                "",  # button_warning_color
                "",  # button_danger_color
                "",  # moysklad_token - теперь глобальный
                "",  # moysklad_stores - теперь глобальный
                False,  # moysklad_enabled - теперь глобальный
                self.shipment_locking_enabled, # shipment_locking_enabled
                article_column_visible,  # article_column_visible
                name_column_visible,     # name_column_visible
                total_qty_column_visible,  # total_qty_column_visible
                stock_column_visible,    # stock_column_visible
                hide_completed_items,    # hide_completed_items
                "",  # cached_server_ip
                True,  # colored_buttons
                ""  # button_colors
            )

            self.logger.debug(
                f"Настройки пользователя {self.current_user} сохранены: "
                f"shipment_columns_width={self.shipment_columns_width}, "
                f"box_columns_width={self.box_columns_width}, "
                f"main_splitter_sizes={self.main_splitter_sizes}, "
                f"article_visible={article_column_visible}, "
                f"name_visible={name_column_visible}, "
                f"total_qty_visible={total_qty_column_visible}, "
                f"stock_visible={stock_column_visible}, "
                f"hide_completed={hide_completed_items}, "
                f"colored_buttons={colored_buttons}"
            )

    def apply_settings(self):
        # Применяем тему
        themes.apply_theme(QApplication.instance(), self.current_theme)

        # Обновляем шрифты
        if self.ui_updater is not None:
            self.ui_updater.update_fonts()

        # Обновляем UI через update_ui, который вызывает update_shipments_tree
        if self.ui_updater is not None:
            self.ui_updater.update_ui()

        # Применяем дополнительные настройки
        self.update_scan_input_style()
        self.apply_button_styles()  # Применяем стили кнопок

        # Применяем стили к чекбоксам после применения темы
        self.apply_toggle_styles()

        # Принудительно обновляем стили таблиц после применения темы
        self._update_table_styles()

        # После обновления UI пересоздаем кнопки "Добавить всё" в таблице поставки
        if (hasattr(self, 'ui_updater') and self.ui_updater is not None and 
            hasattr(self, 'shipment_table') and self.shipment_table):
            # Обновляем стили кнопок в таблице поставки после смены темы
            self.ui_updater._refresh_action_buttons_styles()
            
        # Обновляем видимость кнопки МойСклад
        self.update_moysklad_button_visibility()
            
    def adjust_color_brightness(self, color_str, factor):
        """Изменить яркость цвета на заданный коэффициент"""
        try:
            from PyQt6.QtGui import QColor
            color = QColor(color_str)
            if not color.isValid():
                # Если цвет недействителен, используем зеленый по умолчанию
                color = QColor("#4CAF50")
            # Преобразуем в HSV, изменяем значение (Value), а затем возвращаемся к RGB
            h, s, v, a = color.hsva()
            # Изменяем яркость (v), учитывая ограничения (0-255)
            new_v = min(255, max(0, int(v * factor)))
            new_color = QColor.fromHsv(h, s, new_v, a)
            return new_color.name()
        except:
            # В случае ошибки возвращаем исходный цвет
            return color_str
    
    def apply_button_colors(self):
        """Применить цвета кнопок к верхним кнопкам - всегда цветные"""
        # Цветные кнопки всегда включены
        self._apply_fixed_colorful_button_styles()

    def _apply_fixed_colorful_button_styles(self):
        """Применить фиксированные цветные стили к верхним кнопкам"""
        # Фиксированные цвета для кнопок (по умолчанию)
        default_colors = {
            'new_shipment': "#4CAF50",  # Зеленый
            'new_shipment_gh': "#8BC34A",  # Светло-зеленый
            'new_box': "#2196F3",  # Синий
            'shipment_check': "#9C27B0",  # Фиолетовый
            'print_labels': "#00BCD4",  # Голубой
            'moysklad_sync': "#E91E63",  # Розовый
            'check_stock': "#607D8B"  # Серо-голубой
        }
        
        # Используем сохранённые цвета или цвета по умолчанию
        colors = self.button_colors if self.button_colors else default_colors
        
        # Применяем цвета к кнопкам
        self.new_shipment_btn.setStyleSheet(self.get_button_style(colors.get('new_shipment', default_colors['new_shipment'])))
        self.new_shipment_gh_btn.setStyleSheet(self.get_button_style(colors.get('new_shipment_gh', default_colors['new_shipment_gh'])))
        self.new_box_btn.setStyleSheet(self.get_button_style(colors.get('new_box', default_colors['new_box'])))
        self.shipment_check_btn.setStyleSheet(self.get_button_style(colors.get('shipment_check', default_colors['shipment_check'])))
        self.print_labels_btn.setStyleSheet(self.get_button_style(colors.get('print_labels', default_colors['print_labels'])))
        self.moysklad_sync_btn.setStyleSheet(self.get_button_style(colors.get('moysklad_sync', default_colors['moysklad_sync'])))
        self.check_stock_btn.setStyleSheet(self.get_button_style(colors.get('check_stock', default_colors['check_stock'])))

        self.logger.debug(f"Кнопки перекрашены в случайные цвета: {colors}")
    
    def get_button_style(self, color):
        """Возвращает стиль кнопки с указанным цветом"""
        color_light = self.adjust_color_brightness(color, 1.2)
        color_dark = self.adjust_color_brightness(color, 0.8)
        
        return f"""
            QPushButton {{
                background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                    stop: 0 {color_light},
                    stop: 0.5 {color},
                    stop: 1 {color_dark});
                color: white;
                border: 1px solid {color_dark};
                padding: 4px 10px;
                border-radius: 4px;
                font-weight: 600;
                font-size: 11px;
                min-height: 20px;
                max-height: 20px;
            }}
            QPushButton:hover {{
                background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                    stop: 0 {self.adjust_color_brightness(color, 1.3)},
                    stop: 0.5 {self.adjust_color_brightness(color, 1.1)},
                    stop: 1 {color});
            }}
            QPushButton:pressed {{
                background: qlineargradient(x1: 0, y1: 0, x2: 0, y2: 1,
                    stop: 0 {self.adjust_color_brightness(color, 0.9)},
                    stop: 0.5 {self.adjust_color_brightness(color, 0.8)},
                    stop: 1 {self.adjust_color_brightness(color, 0.7)});
            }}
        """
    
    def on_recolor_buttons_clicked(self):
        """Обработкатчик нажатия кнопки перекраски кнопок"""
        self._apply_random_colorful_button_styles()
        self.logger.info("Кнопки перекрашены через настройки")

    def _apply_random_colorful_button_styles(self):
        """Применить случайные цветные стили к верхним кнопкам"""
        from dialogs import generate_random_color
        
        # Генерируем случайные цвета для каждой кнопки
        colors = {
            'new_shipment': generate_random_color(),
            'new_shipment_gh': generate_random_color(),
            'new_box': generate_random_color(),
            'shipment_check': generate_random_color(),
            'print_labels': generate_random_color(),
            'moysklad_sync': generate_random_color(),
            'check_stock': generate_random_color()
        }
        
        # Сохраняем цвета
        self.button_colors = colors
        
        # Применяем цвета к кнопкам
        self.new_shipment_btn.setStyleSheet(self.get_button_style(colors['new_shipment']))
        self.new_shipment_gh_btn.setStyleSheet(self.get_button_style(colors['new_shipment_gh']))
        self.new_box_btn.setStyleSheet(self.get_button_style(colors['new_box']))
        self.shipment_check_btn.setStyleSheet(self.get_button_style(colors['shipment_check']))
        self.print_labels_btn.setStyleSheet(self.get_button_style(colors['print_labels']))
        self.moysklad_sync_btn.setStyleSheet(self.get_button_style(colors['moysklad_sync']))
        self.check_stock_btn.setStyleSheet(self.get_button_style(colors['check_stock']))
        
        self.logger.info(f"Кнопки перекрашены в случайные цвета: {colors}")

    def open_settings(self):
        """Открыть диалог настроек"""
        from dialogs import SettingsDialog
        
        dialog = SettingsDialog(
            self.font_size, self.label_font_size, self.current_theme,
            self.ok_sound, self.error_sound,
            True, self.tone_sound, self.sound_volume, self
        )

        # Подключаем кнопку перекраски кнопок
        dialog.recolor_buttons_btn.clicked.connect(self.on_recolor_buttons_clicked)

        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.font_size = dialog.get_font_size()
            self.label_font_size = dialog.get_label_font_size()
            self.current_theme = dialog.get_theme()
            self.ok_sound = dialog.get_ok_sound()
            self.error_sound = dialog.get_error_sound()
            self.tone_sound = dialog.get_tone_sound()
            self.sound_volume = dialog.get_sound_volume()
            
            # Применяем громкость звуков
            utils.set_sound_volume(self.sound_volume)

            # Сохраняем все настройки одним вызовом
            if self.current_user:
                # Получаем текущующие значения чекбоксов
                article_visible = self.article_display_checkbox.isChecked() if hasattr(self, 'article_display_checkbox') else True
                name_visible = self.name_display_checkbox.isChecked() if hasattr(self, 'name_display_checkbox') else False
                total_qty_visible = self.total_qty_display_checkbox.isChecked() if hasattr(self, 'total_qty_display_checkbox') else True
                stock_visible = self.stock_display_checkbox.isChecked() if hasattr(self, 'stock_display_checkbox') else True
                hide_completed = self.hide_completed_checkbox.isChecked() if hasattr(self, 'hide_completed_checkbox') else False

                self.logger.info(f"Сохранение настроек видимости: article={article_visible}, name={name_visible}, total_qty={total_qty_visible}, stock={stock_visible}, hide_completed={hide_completed}")

                import json
                button_colors_json = json.dumps(self.button_colors)
                self.logger.info(f"Сохранение цветов кнопок: {button_colors_json}")

                # Сохраняем настройки пользователя (настройки МойСклад теперь сохраняются в диалоге)
                database.set_user_settings(
                    self.current_user, self.font_size, self.label_font_size, self.current_theme,
                    self.ok_sound, self.error_sound,
                    self.shipment_columns_width, self.box_columns_width, self.main_splitter_sizes,
                    self.window_width, self.window_height, "",
                    "", "", "",
                    "", "", False,  # moysklad_token, moysklad_stores, moysklad_enabled - теперь глобальные
                    False,  # shipment_locking_enabled
                    article_visible, name_visible, total_qty_visible, stock_visible, hide_completed,
                    "",  # cached_server_ip
                    True,  # colored_buttons
                    button_colors_json,  # button_colors
                    self.tone_sound,  # tone_sound
                    self.sound_volume  # sound_volume
                )

                self.logger.info("Настройки пользователя сохранены")

            # Применяем настройки
            self.apply_settings()
            self.apply_button_styles()
            self.update_user_display()

    def open_database_settings(self):
        """Открыть настройки базы данных (теперь вызывается из вкладки настроек)"""
        from db_settings_dialog import DatabaseSettingsDialog
        dialog = DatabaseSettingsDialog(self)
        dialog.exec()

    def save_session(self):
        if self.saving:
            return
        self.saving = True
        try:
            if self.current_shipment:
                # ОПТИМИЗАЦИЯ: Инкрементальное сохранение только текущей коробки
                # Это значительно быстрее при работе с удалённой БД
                self.shipment_manager._save_current_box_incremental()
        finally:
            self.saving = False
    
    def save_all_shipments(self):
        """Сохраняет ВСЕ поставки в базу данных (при закрытии программы)"""
        try:
            # Отменяем любые отложенные сохранения
            if hasattr(self.shipment_manager, 'save_timer') and self.shipment_manager.save_timer:
                self.shipment_manager.save_timer.stop()
                self.shipment_manager.save_timer = None
            self.shipment_manager.save_pending = False

            # Сохраняем каждую поставку с её коробками
            saved_count = 0
            for shipment in self.shipments.values():
                try:
                    self.data_controller.save_shipment(shipment)
                    saved_count += 1
                except Exception as e:
                    self.logger.warning(f"Не удалось сохранить поставку {shipment.destination_name}: {e}")
            
            self.logger.info(f"Сохранено {saved_count} поставок при закрытии программы")
        except Exception as e:
            self.logger.error(f"Ошибка при сохранении всех поставок: {e}", exc_info=True)

    def force_save_session(self):
        """Force immediate save without checking for pending saves"""
        try:
            # If there are pending saves in the shipment manager, perform them now
            if hasattr(self.shipment_manager, 'perform_save') and self.shipment_manager.save_pending:
                self.shipment_manager.perform_save()

            # Also cancel any pending scheduled saves
            if hasattr(self.shipment_manager, 'save_timer') and self.shipment_manager.save_timer:
                self.shipment_manager.save_timer.stop()
                self.shipment_manager.save_timer = None
                self.shipment_manager.save_pending = False

            # Then perform the normal save
            self.save_session()

            # Force database commit to ensure data is written
            # ИСПОЛЬЗУЕМ execute_query для коммита без закрытия соединения
            from db_connection import execute_query
            execute_query("SELECT 1")  # Просто чтобы убедиться, что соединение живо
        except Exception as e:
            self.logger.error(f"Ошибка при принудительном сохранении сессии: {e}", exc_info=True)

    def load_all_data(self):
        """Загрузка всех данных из базы данных (асинхронно)"""
        self.show_progress("Загрузка данных...", 100)
        self.logger.info("Запуск асинхронной загрузки данных из базы данных")

        # Загружаем данные в фоновом потоке
        self.async_manager.execute_async(
            self.data_controller.load_shipments,
            callback=self._on_load_shipments_finished,
            error_callback=self._on_load_shipments_error
        )

    def _on_load_shipments_finished(self, data):
        """Обработка успешной загрузки данных (вызывается в главном потоке)"""
        try:
            self.shipments = data['shipments']
            self.group_shipments = data['group_shipments']

            self.logger.info(f"Загружено {len(self.shipments)} поставок и {len(self.group_shipments)} групп")

            # Восстанавливаем текущую поставку
            if self.current_shipment:
                found_shipment = None
                if self.current_shipment.destination_name in self.shipments:
                    found_shipment = self.shipments[self.current_shipment.destination_name]
                else:
                    for group_shipment in self.group_shipments.values():
                        for sub_shipment_key, sub_shipment in group_shipment.sub_shipments.items():
                            if (hasattr(sub_shipment, 'original_destination_name') and
                                sub_shipment.original_destination_name == self.current_shipment.destination_name) or \
                               (sub_shipment.destination_name == self.current_shipment.destination_name):
                                found_shipment = sub_shipment
                                break
                        if found_shipment:
                            break

                self.current_shipment = found_shipment
            else:
                self.current_shipment = None

            # Инициализируем кэш
            for shipment in self.shipments.values():
                self.shipment_manager.update_cache(shipment)

            if self.first_load:
                self.first_load = False

            # Обновляем сессию пользователя
            if self.current_shipment and hasattr(self, 'current_user') and self.current_user:
                from database import update_user_session
                shipment_name_for_session = getattr(self.current_shipment, 'original_destination_name', self.current_shipment.destination_name)
                update_user_session(shipment_name_for_session, self.current_user)

            self.hide_progress("Данные загружены", 2000)
            self.logger.info("Загрузка всех данных завершена")

            # Обновляем UI с загруженными данными
            if self.ui_updater is not None:
                self.ui_updater.update_ui()
        except Exception as e:
            self.logger.error(f"Ошибка при обработке загруженных данных: {e}", exc_info=True)
            self.hide_progress("Ошибка загрузки данных", 3000)
            QMessageBox.warning(self, "Ошибка", f"Не удалось обработать данные: {e}")

    def _on_load_shipments_error(self, error_msg):
        """Обработка ошибки загрузки данных"""
        self.logger.error(f"Ошибка при загрузке данных: {error_msg}")
        self.hide_progress("Ошибка загрузки данных", 3000)
        QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить данные: {error_msg}")

    def update_ui(self):
        # Обновляем UI без проверки флага при явном вызове (например, при выборе поставки/коробки)
        # Устанавливаем флаг обновления, чтобы избежать рекурсивных вызовов
        if self.updating_ui:
            return  # Избегаем рекурсивных вызовов

        self.updating_ui = True
        try:
            # Проверяем, инициализирован ли ui_updater
            if self.ui_updater is not None:
                self.ui_updater.update_ui()
        except Exception as e:
            self.logger.error(f"Ошибка при обновлении UI: {e}", exc_info=True)
        finally:
            self.updating_ui = False
            
    def start_new_shipment(self):
        """Создание новой поставки (обычной или группуовой)
        Тип поставки определяется автоматически по структуре файла Excel"""
        self.shipment_operations.start_new_shipment()

    def update_group_shipment_from_google_sheets(self, group_shipment):
        """Обновление групповой поставки из Google Sheets"""
        try:
            import gspread
            from google.oauth2.service_account import Credentials
        except ImportError:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "Ошибка", "Установите gspread: pip install gspread google-auth")
            return
        
        import os
        credentials_path = os.path.join(os.path.dirname(__file__), "e-object-470910-p6-3500f3ddbdd3.json")
        if not os.path.exists(credentials_path):
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Ошибка", f"Файл credentials не найден: {credentials_path}")
            return
        
        try:
            scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
            creds = Credentials.from_service_account_file(credentials_path, scopes=scope)
            client = gspread.authorize(creds)
            
            # ID таблицы поставок
            spreadsheet_id = "1OGgsS0T4qaEekJgEkVTplZfoeQ7MeMth8o8eJTqnJGA"
            spreadsheet = client.open_by_key(spreadsheet_id)
            
            # Получаем список листов
            worksheets = spreadsheet.worksheets()
            if not worksheets:
                from PyQt6.QtWidgets import QMessageBox
                QMessageBox.warning(self, "Ошибка", "Таблица не содержит листов")
                return
            
            sheet_names = [ws.title for ws in worksheets]
            
            # Показываем диалог выбора листа
            dialog = GoogleSheetsUpdateDialog(sheet_names, group_shipment.group_name, self)
            if dialog.exec() != QDialog.DialogCode.Accepted:
                return
            
            sheet_name = dialog.get_sheet_name()
            
            # Запускаем асинхронное обновление
            self.show_busy_progress(f"Обновление группы '{sheet_name}'...")
            self.logger.info(f"Запуск обновления групповой поставки из Google Sheets: {sheet_name}")
            
            self.async_manager.execute_async(
                self._update_group_shipment_from_sheet_worker,
                callback=self._on_update_group_shipment_finished,
                error_callback=self._on_update_group_shipment_error,
                spreadsheet=spreadsheet,
                sheet_name=sheet_name,
                group_shipment=group_shipment
            )
        except Exception as e:
            self.logger.error(f"Ошибка при обновлении из Google Sheets: {e}", exc_info=True)
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Ошибка", f"Ошибка при обновлении из Google Sheets:\n{e}")

    def _update_group_shipment_from_sheet_worker(self, spreadsheet, sheet_name, group_shipment):
        """Worker для обновления групповой поставки из Google Sheets"""
        import pandas as pd
        
        worksheet = spreadsheet.worksheet(sheet_name)
        all_values = worksheet.get_all_values()
        
        if not all_values or len(all_values) < 2:
            raise ValueError(f"Лист '{sheet_name}' пуст")
        
        # Находим строку заголовков
        header_row_idx = 0
        for i, row in enumerate(all_values[:10]):
            row_lower = [str(cell).lower() for cell in row]
            if any('штрихкод' in cell or 'шк' in cell or 'barcode' in cell for cell in row_lower):
                header_row_idx = i
                break
        
        headers = all_values[header_row_idx]
        data_rows = all_values[header_row_idx + 1:]
        
        if not data_rows:
            raise ValueError(f"Лист '{sheet_name}' не содержит данных")
        
        df = pd.DataFrame(data_rows, columns=headers)
        
        # Определяем колонку штрихкодов
        barcode_col = None
        for col in df.columns:
            col_lower = str(col).lower().strip()
            if 'штрихкод' in col_lower or 'шк' in col_lower or 'barcode' in col_lower:
                barcode_col = col
                break
        
        if barcode_col is None:
            barcode_col = df.columns[0]
        
        return {
            'df': df,
            'barcode_col': barcode_col,
            'sheet_name': sheet_name,
            'group_shipment': group_shipment
        }

    def _on_update_group_shipment_finished(self, result):
        """Обработка успешного обновления групповой поставки"""
        self.hide_progress(f"Группа '{result['sheet_name']}' обновлена", 3000)
        self.shipment_operations.update_group_shipment_from_google_sheets_data(result)

    def _on_update_group_shipment_error(self, error_msg):
        """Обработка ошибки обновления"""
        self.hide_progress("Ошибка обновления", 3000)
        self.logger.error(f"Ошибка обновления из Google Sheets: {error_msg}")
        from PyQt6.QtWidgets import QMessageBox
        QMessageBox.critical(self, "Ошибка", f"Ошибка обновления из Google Sheets:\n{error_msg}")

    def import_shipment_from_google_sheets(self):
        """Импорт поставки из Google Sheets"""
        try:
            import gspread
            from google.oauth2.service_account import Credentials
        except ImportError:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "Ошибка", "Установите gspread: pip install gspread google-auth")
            return
        
        import os
        credentials_path = os.path.join(os.path.dirname(__file__), "e-object-470910-p6-3500f3ddbdd3.json")
        if not os.path.exists(credentials_path):
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Ошибка", f"Файл credentials не найден: {credentials_path}")
            return
        
        try:
            scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
            creds = Credentials.from_service_account_file(credentials_path, scopes=scope)
            client = gspread.authorize(creds)
            
            # ID таблицы поставок
            spreadsheet_id = "1OGgsS0T4qaEekJgEkVTplZfoeQ7MeMth8o8eJTqnJGA"
            spreadsheet = client.open_by_key(spreadsheet_id)
            
            # Получаем список листов
            worksheets = spreadsheet.worksheets()
            if not worksheets:
                from PyQt6.QtWidgets import QMessageBox
                QMessageBox.warning(self, "Ошибка", "Таблица не содержит листов")
                return
            
            sheet_names = [ws.title for ws in worksheets]
            
            # Показываем диалог выбора листа с опцией групповой поставки
            dialog = GoogleSheetsImportDialog(sheet_names, self)
            if dialog.exec() != QDialog.DialogCode.Accepted:
                return
            
            sheet_name = dialog.get_sheet_name()
            is_group = dialog.is_group_shipment()
            
            # Запускаем асинхронный импорт
            self.show_busy_progress(f"Импорт поставки '{sheet_name}'...")
            self.logger.info(f"Запуск импорта поставки из Google Sheets: {sheet_name} (групповая={is_group})")
            
            self.async_manager.execute_async(
                self._import_shipment_from_sheet_worker,
                callback=self._on_import_shipment_finished,
                error_callback=self._on_import_shipment_error,
                spreadsheet=spreadsheet,
                sheet_name=sheet_name,
                is_group=is_group
            )
        except Exception as e:
            self.logger.error(f"Ошибка при импорте из Google Sheets: {e}", exc_info=True)
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Ошибка", f"Ошибка при импорте из Google Sheets:\n{e}")

    def _import_shipment_from_sheet_worker(self, spreadsheet, sheet_name, is_group=None):
        """Worker для импорта поставки из Google Sheets в фоновом потоке"""
        import pandas as pd
        from io import StringIO
        
        worksheet = spreadsheet.worksheet(sheet_name)
        
        # Получаем все данные из листа
        all_values = worksheet.get_all_values()
        
        if not all_values or len(all_values) < 2:
            raise ValueError(f"Лист '{sheet_name}' пуст или содержит только заголовки")
        
        # Находим строку заголовков
        header_row_idx = 0
        for i, row in enumerate(all_values[:10]):
            row_lower = [str(cell).lower() for cell in row]
            if any('штрихкод' in cell or 'шк' in cell or 'barcode' in cell for cell in row_lower):
                header_row_idx = i
                break
        
        # Создаём DataFrame
        headers = all_values[header_row_idx]
        data_rows = all_values[header_row_idx + 1:]
        
        if not data_rows:
            raise ValueError(f"Лист '{sheet_name}' не содержит данных")
        
        # Преобразуем в формат для pandas
        df = pd.DataFrame(data_rows, columns=headers)
        
        # Определяем колонку штрихкодов
        barcode_col = None
        for col in df.columns:
            col_lower = str(col).lower().strip()
            if 'штрихкод' in col_lower or 'шк' in col_lower or 'barcode' in col_lower:
                barcode_col = col
                break
        
        if barcode_col is None and len(df.columns) > 0:
            barcode_col = df.columns[0]
        
        if barcode_col is None:
            raise ValueError("Не найдена колонка со штрихкодами")
        
        # Если пользователь явно указал тип поставки - используем это
        if is_group is not None:
            pass  # Используем значение от пользователя
        else:
            # Авто-определение (для обратной совместимости)
            significant_columns = [col for col in df.columns if str(col).strip() and not str(col).startswith('Unnamed')]
            quantity_cols = [col for col in significant_columns 
                            if any(kw in str(col).lower() for kw in ['количество', 'кол-во', 'qty', 'quantity'])]
            
            is_group = len(quantity_cols) > 1 or len(significant_columns) > 4
            
            # Дополнительная проверка: считаем колонки с числовыми данными после первых 2-3 колонок
            if not is_group and len(significant_columns) > 3:
                numeric_cols = 0
                for col in significant_columns[2:]:
                    sample_val = df[col].dropna().iloc[0] if len(df[col].dropna()) > 0 else None
                    if sample_val is not None:
                        try:
                            float(str(sample_val).replace(',', '.').strip())
                            numeric_cols += 1
                        except (ValueError, TypeError):
                            pass
                if numeric_cols > 1:
                    is_group = True
                    self.logger.info(f"Обнаружена групповая поставка: {numeric_cols} колонок с числами")
        
        return {
            'sheet_name': sheet_name,
            'df': df,
            'barcode_col': barcode_col,
            'is_group': is_group,
        }

    def _on_import_shipment_finished(self, result):
        """Обработка успешного импорта поставки из Google Sheets"""
        self.hide_progress(f"Поставка '{result['sheet_name']}' импортирована", 3000)
        self.logger.info(f"Импорт поставки из Google Sheets завершён: {result['sheet_name']}")
        
        # Передаём данные в shipment_operations для создания поставки
        self.shipment_operations.create_shipment_from_google_sheets_data(result)

    def _on_import_shipment_error(self, error_msg):
        """Обработка ошибки импорта поставки из Google Sheets"""
        self.hide_progress("Ошибка импорта", 3000)
        self.logger.error(f"Ошибка импорта из Google Sheets: {error_msg}")
        from PyQt6.QtWidgets import QMessageBox
        QMessageBox.critical(self, "Ошибка", f"Ошибка импорта из Google Sheets:\n{error_msg}")

    def new_box(self):
        self.shipment_manager.new_box()

    def handle_scan(self):
        self.shipment_manager.handle_scan()

    def add_all_remaining_to_box_by_barcode(self, barcode, qty=None):
        """Добавить указанное количество товара в коробкуку по штрихкоду
        
        Args:
            barcode: штрихкод товара
            qty: количество для добавления (по умолчанию = remaining_qty)
        """
        self.shipment_manager.add_all_remaining_to_box_by_barcode(barcode, qty)

    def update_shipment_composition(self):
        self.shipment_operations.update_shipment_composition()

    def show_shipment_table_context_menu(self, position):
        """Отображение контекстного меню для таблицы поставки"""
        if not self.current_shipment:
            return
            
        item = self.shipment_table.itemAt(position)
        if not item:
            return
            
        row = item.row()
        barcode_item = self.shipment_table.item(row, 0)
        if not barcode_item:
            return
            
        barcode = barcode_item.text()
        
        # Проверяем, что товар существует в поставке
        if barcode not in self.current_shipment.shipment_items:
            return
            
        shipment_item = self.current_shipment.shipment_items[barcode]
        
        menu = QMenu(self)
        
        # Добавляем действие для изменения количества
        edit_qty_action = menu.addAction("Изменить количество")
        
        # Добавляем действие для удаления позиции
        delete_action = menu.addAction("Удалить из поставки")
        
        # Добавляем действие для добавления новой позиции
        add_action = menu.addAction("Добавить новую позицию")
        
        action = menu.exec(self.shipment_table.mapToGlobal(position))
        
        if action == edit_qty_action:
            # Открываем диалог для изменения количества
            self.open_shipment_quantity_edit_dialog(row, shipment_item)
        elif action == delete_action:
            # Удаляем позицию из поставки
            self.current_shipment.remove_shipment_item(barcode)
            self.current_shipment.invalidate_caches()
            self.ui_updater.update_current_components()
            # Сохраняем изменения в БД
            self.data_controller.save_shipment(self.current_shipment)
        elif action == add_action:
            # Добавляем новую позицию
            self.add_new_shipment_item()
    
    def open_shipment_quantity_edit_dialog(self, row, shipment_item):
        """Открытие диалога для изменения количества товара в поставке"""
        # Получаем текущующее количество
        current_qty = shipment_item.total_qty
        
        # Создаем и показываем диалог ввода
        from PyQt6.QtWidgets import QInputDialog
        new_qty_str, ok = QInputDialog.getText(
            self, 
            "Изменить количество", 
            f"Введите новое количество для товара {shipment_item.barcode} ({shipment_item.sku}):",
            text=str(current_qty)
        )
        
        if ok and new_qty_str:
            try:
                new_qty = int(new_qty_str)
                if new_qty < 0:
                    QMessageBox.warning(self, "Ошибка", "Количество не может быть отрицательным!")
                    return
                
                # Проверяем, не превышает ли новое количество распределенное количество
                if new_qty < shipment_item.allocated_qty:
                    reply = QMessageBox.question(
                        self, "Подтверждение",
                        f"Новое количество ({new_qty}) меньше распределенного количества ({shipment_item.allocated_qty}).\n"
                        f"Это приведет к перемещению {shipment_item.allocated_qty - new_qty} шт. в список удаленных.\n"
                        "Продолжить?",
                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                    )
                    if reply != QMessageBox.StandardButton.Yes:
                        return
                
                # Обновляем количество
                old_qty = shipment_item.total_qty
                shipment_item.total_qty = new_qty
                
                # Обновляем removed_items при необходимости
                qty_diff = new_qty - old_qty
                if qty_diff < 0 and shipment_item.barcode in self.current_shipment.removed_items:
                    # Уменьшаем количество в removed_items если товар уже удален частично
                    excess_reduction = abs(qty_diff)
                    removed_data = self.current_shipment.removed_items[shipment_item.barcode]
                    new_removed_qty = max(0, removed_data['allocated_qty'] - excess_reduction)
                    if new_removed_qty <= 0:
                        del self.current_shipment.removed_items[shipment_item.barcode]
                    else:
                        removed_data['allocated_qty'] = new_removed_qty
                elif qty_diff < 0:
                    # Если уменьшаем количество и это приводит к превышению распределения
                    excess_qty = shipment_item.allocated_qty - new_qty
                    if excess_qty > 0:
                        self.current_shipment.removed_items[shipment_item.barcode] = {
                            'sku': shipment_item.sku,
                            'allocated_qty': excess_qty
                        }
                
                # Сбрасываем кэши в модели поставки при изменении количества товара
                self.current_shipment.invalidate_caches()

                # Обновляем интерфейс
                self.ui_updater.update_current_components()
                self.ui_updater.update_shipments_tree()
                
                # Статус скрыт, сообщения не отображаются
                # self.statusBar().showMessage(f"Количество товара {shipment_item.barcode} изменено с {old_qty} на {new_qty}", 3000)
                
                # Структурное изменение — полное сохранение
                if hasattr(self.shipment_manager, 'schedule_full_save'):
                    self.shipment_manager.schedule_full_save()
                elif hasattr(self.shipment_manager, 'schedule_save'):
                    self.shipment_manager.schedule_save()
            except ValueError:
                QMessageBox.warning(self, "Ошибка", "Введите корректное число!")
    
    def add_new_shipment_item(self):
        """Добавление новой позиции в ппоставккуавкиу"""
        # Создаем диалог для ввода данных новой позиции
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton, QSpinBox
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Добавить новую позицию")
        dialog.resize(400, 150)
        
        layout = QVBoxLayout()
        
        # Поле для штрихкода
        barcode_layout = QHBoxLayout()
        barcode_layout.addWidget(QLabel("Штрихкод:"))
        barcode_input = QLineEdit()
        barcode_layout.addWidget(barcode_input)
        layout.addLayout(barcode_layout)
        
        # Поле для артикула
        sku_layout = QHBoxLayout()
        sku_layout.addWidget(QLabel("Артикул:"))
        sku_input = QLineEdit()
        sku_layout.addWidget(sku_input)
        layout.addLayout(sku_layout)
        
        # Поле для количества
        qty_layout = QHBoxLayout()
        qty_layout.addWidget(QLabel("Количество:"))
        qty_input = QSpinBox()
        qty_input.setRange(1, 9999)
        qty_input.setValue(1)
        qty_layout.addWidget(qty_input)
        layout.addLayout(qty_layout)
        
        # Кнопки OK и Cancel
        buttons_layout = QHBoxLayout()
        ok_button = QPushButton("OK")
        cancel_button = QPushButton("Отмена")
        
        def on_ok():
            barcode = barcode_input.text().strip()
            sku = sku_input.text().strip()
            qty = qty_input.value()
            
            if not barcode:
                QMessageBox.warning(dialog, "Ошибка", "Введите штрихкод!")
                return
            if not sku:
                QMessageBox.warning(dialog, "Ошибка", "Введите артикул!")
                return
            if qty <= 0:
                QMessageBox.warning(dialog, "Ошибка", "Количество должно быть большеше 0!")
                return
                
            # Проверяем, существует ли уже ткакой штрихкод
            if barcode in self.current_shipment.shipment_items:
                QMessageBox.warning(dialog, "Ошибка", "Товар с ткаким штрихкодом уже существует в поставке!")
                return
                
            # Добавляем новый товар в ппоставккуавкиу
            self.current_shipment.add_shipment_item(barcode, sku, qty)
            
            # Сбрасываем кэши в модели поставки при добавлении товара
            self.current_shipment.invalidate_caches()

            # Обновляем интерфейс
            self.ui_updater.update_current_components()
            self.ui_updater.update_shipments_tree()
            
            # Статус скрыт, сообщения не отображаются
            # self.statusBar().showMessage(f"Добавлен товар {barcode} ({sku}) в количестве {qty}", 3000)
            
            # Отложенное сохранение
            if hasattr(self.shipment_manager, 'schedule_save'):
                self.shipment_manager.schedule_save()
                
            dialog.accept()
            
        ok_button.clicked.connect(on_ok)
        cancel_button.clicked.connect(dialog.reject)
        
        buttons_layout.addWidget(ok_button)
        buttons_layout.addWidget(cancel_button)
        
        layout.addLayout(buttons_layout)
        dialog.setLayout(layout)
        
        dialog.exec()

    def on_shipment_cell_changed(self, row, column):
        self.shipment_manager.on_shipment_cell_changed(row, column)

    def on_box_cell_changed(self, row, column):
        self.shipment_manager.on_box_cell_changed(row, column)

    def on_shipment_clicked(self, item):
        """Обработка одинарного клика по поставке или коробке в дереве"""
        if not item:
            return

        if self.updating_ui:
            return

        shipment = item.data(0, Qt.ItemDataRole.UserRole)
        box = item.data(0, Qt.ItemDataRole.UserRole + 1)
        group_shipment = item.data(0, Qt.ItemDataRole.UserRole + 2)

        if group_shipment and not shipment and not box:  # Это групповая поставка
            self.current_shipment = None
            self.expand_current_shipment_collapse_others(group_shipment)
            
            # Скрываем центральную панель полностью
            center_widget = self.main_splitter.widget(1)  # Центральный виджет
            if center_widget:
                center_widget.hide()
            
            self.ui_updater.update_group_shipment_summary(group_shipment)
            self.ui_updater.update_group_shipment_boxes_table(group_shipment)
            self.ui_updater.update_group_shipment_items_table(group_shipment)
            self.ui_updater.update_shipment_tree_selection()
            return

        if shipment and not box:  # Это поставка
            self.current_shipment = shipment
            shipment.current_box_index = -1

            self.expand_current_shipment_collapse_others(shipment)
            
            # Показываем центральную панель
            center_widget = self.main_splitter.widget(1)  # Центральный виджет
            if center_widget:
                center_widget.show()

            hide_completed = getattr(shipment, 'hide_completed_items', self.hide_completed_items_setting)
            self.hide_completed_checkbox.blockSignals(True)
            self.hide_completed_checkbox.setChecked(hide_completed)
            self.hide_completed_checkbox.blockSignals(False)
            shipment.hide_completed_items = hide_completed
            self.hide_completed_items_setting = hide_completed

            self.ui_updater.reset_group_shipment_view()
            self.ui_updater.update_current_box_label()
            self.ui_updater.update_current_box_table()
            self.ui_updater.update_shipment_tree_selection()
            self.ui_updater.update_shipment_table_rows_visibility()

            from PyQt6.QtCore import QTimer
            QTimer.singleShot(50, self.update_user_activity)
            QTimer.singleShot(150, lambda: self._delayed_box_switch_update(False))

        elif shipment and box:  # Это коробка
            self.current_shipment = shipment

            box_found = False
            for i, b in enumerate(shipment.boxes):
                if b == box:
                    shipment.current_box_index = i
                    box_found = True
                    break

            if not box_found:
                logger.error(f"Коробка {box.box_id} не найдена в списке коробок поставки {shipment.destination_name}")
                shipment.current_box_index = -1

            self.expand_current_shipment_collapse_others(shipment)
            
            # Показываем центральную панель
            center_widget = self.main_splitter.widget(1)  # Центральный виджет
            if center_widget:
                center_widget.show()

            hide_completed = getattr(shipment, 'hide_completed_items', self.hide_completed_items_setting)
            self.hide_completed_checkbox.blockSignals(True)
            self.hide_completed_checkbox.setChecked(hide_completed)
            self.hide_completed_checkbox.blockSignals(False)
            shipment.hide_completed_items = hide_completed
            self.hide_completed_items_setting = hide_completed

            self.ui_updater.reset_group_shipment_view()
            self.ui_updater.update_current_box_label()
            self.ui_updater.update_current_box_table()
            self.ui_updater.update_shipment_tree_selection()
            self.ui_updater.update_shipment_table_rows_visibility()
            self.apply_button_colors()

            from PyQt6.QtCore import QTimer
            QTimer.singleShot(50, self.update_user_activity)
            QTimer.singleShot(150, lambda: self._delayed_box_switch_update(is_group=False))

    def _delayed_box_switch_update(self, is_group=False):
        """Отложенное обновление при переключении коробки"""
        try:
            current = self.shipments_tree_widget.currentItem()
            if current:
                group_shipment = current.data(0, Qt.ItemDataRole.UserRole + 2)
                if group_shipment:
                    self.ui_updater.update_group_shipment_items_table(group_shipment)
                    self.ui_updater._update_shipments_tree_progress()
                    self.ui_updater._refresh_action_buttons_styles()
                    return

            self.ui_updater.update_shipment_table()
            self.ui_updater._update_shipments_tree_progress()
            self.ui_updater._refresh_action_buttons_styles()
            self._update_add_all_button_visibility()
        except Exception:
            pass

    _delayed_shipment_switch_update = _delayed_box_switch_update  # Alias для поставки

    def on_shipment_double_clicked(self, item):
        """Обработкатка двойного клика по поставке или коробкуке в дереве"""
        import logging
        logger = logging.getLogger(__name__)
        logger.info("Начало обработки двойного клика по элементу дерева")

        if not item:
            logger.warning("Получен клик по пустому элементу")
            return

        # Защита от рекурсивных вызовов
        if self.updating_ui:
            logger.debug("on_shipment_double_clicked: пропущено из-за updating_ui=True")
            return

        # Получаем данные из пользовательских ролей
        shipment = item.data(0, Qt.ItemDataRole.UserRole)
        box = item.data(0, Qt.ItemDataRole.UserRole + 1)
        
        # Сохраняем предыдущую поставку перед переключением
        if self.current_shipment and self.current_shipment != shipment:
            try:
                self.data_controller.save_shipment(self.current_shipment)
            except Exception:
                pass
        
        if shipment and not box:  # Это поставка
            logger.info(f"Выбрана поставка: {shipment.destination_name}, обновление UI нача��о")
            
            self.current_shipment = shipment

            # Сбрасываем кэш прогресса для корректного отображения
            shipment.invalidate_caches()


            # Сбрасываем индекс текущей коробки при выборе поставки
            shipment.current_box_index = -1
            
            self.expand_current_shipment_collapse_others(shipment)
            
            # Update UI immediately for responsiveness, but delay database save
            # Используем новую реализацию обновления UI для большешей эффективности
            self.ui_updater.update_current_components(full_update=True)
            self.ui_updater.update_shipments_tree()

            # Статус скрыт, сообщения не отображаются
            # self.statusBar().showMessage(f"Выбрана поставка: {shipment.destination_name}", 3000)
            
            # Обновляем активность пользователя в выбранной поставке
            self.update_user_activity()
            
            # Schedule a delayed save to the database to avoid performance issues
            if hasattr(self.shipment_manager, 'schedule_save'):
                self.shipment_manager.schedule_save()
      
        elif shipment and box:  # Это коробкука
            self.current_shipment = shipment

            # Сбрасываем кэш прогресса для корректного отображения
            shipment.invalidate_caches()


            # Находим индекс коробки в списке коробкуок поставки
            box_found = False
            for i, b in enumerate(shipment.boxes):
                if b == box:
                    shipment.current_box_index = i
                    box_found = True
                    break
            
            # Если коробкука не найдена, выводим сообщение об ошибке
            if not box_found:
                logger.error(f"��оробка {box.box_id} не найдена в списке коробкуок поставки {shipment.destination_name}")
                shipment.current_box_index = -1
            
            self.expand_current_shipment_collapse_others(shipment)
            # Update UI immediately for responsiveness, but delay database save
            # Используем новую реализацию обновления UI для большешей эффективности
            self.ui_updater.update_current_components(full_update=True)
            self.ui_updater.update_shipments_tree()

            # Применяем цвета кнопок заново, чтобы убедиться, что они не сбрасываются
            self.apply_button_colors()
            # Статус скрыт, сообщения не отображаются
            # self.statusBar().showMessage(f"Выбрана коробка: {box.box_id} в поставке {shipment.destination_name}", 300)
            
            # Обновляем активность пользователя в выбранной поставке
            self.update_user_activity()
            
            # Schedule a delayed save to the database to avoid performance issues
            if hasattr(self.shipment_manager, 'schedule_save'):
                self.shipment_manager.schedule_save()
            
    def update_user_activity(self):
        """Обновление кактивности пользователя в текущей поставке"""
        if (self.current_shipment and
            hasattr(self, 'current_user') and
            self.current_user):
            try:
                from database import update_user_session
                # Используем имя поставки для обновления сессии в базе данных
                shipment_name_for_session = self.current_shipment.destination_name
                update_user_session(shipment_name_for_session, self.current_user)
            except Exception as e:
                self.logger.error(f"Ошибка при обновлении сессии пользователя: {e}", exc_info=True)
      
    def expand_current_shipment_collapse_others(self, current_shipment):
        """Разворачивает текущую поставку и сворачивает все остальные"""
        # Быстро обновляем флаги без лишних операций
        for shipment_name, shipment in self.shipments.items():
            shipment.is_expanded = (shipment == current_shipment)

        for group_name, group_shipment in self.group_shipments.items():
            group_shipment.is_expanded = False
            for shipment_name, shipment in group_shipment.sub_shipments.items():
                if shipment == current_shipment:
                    shipment.is_expanded = True
                    group_shipment.is_expanded = True
                else:
                    shipment.is_expanded = False

        # Не вызываем update_shipment_tree_expansion() - флаги будут применены в update_shipment_tree_selection()

    def on_shipment_expanded(self, item):
        """Обработкатка разворачивания поставки или группуы"""
        shipment = item.data(0, Qt.ItemDataRole.UserRole)
        box = item.data(0, Qt.ItemDataRole.UserRole + 1)
        group_shipment = item.data(0, Qt.ItemDataRole.UserRole + 2)
        
        if shipment and not box and not group_shipment:
            shipment.is_expanded = True
            # Schedule a delayed save to the database to avoid performance issues
            if hasattr(self.shipment_manager, 'schedule_save'):
                self.shipment_manager.schedule_save()
            
        elif group_shipment:
            group_shipment.is_expanded = True
            # Schedule a delayed save to the database to avoid performance issues
            if hasattr(self.shipment_manager, 'schedule_save'):
                self.shipment_manager.schedule_save()

    def on_shipment_collapsed(self, item):
        """Обработкатка сворачивания поставки ��ли группуы"""
        shipment = item.data(0, Qt.ItemDataRole.UserRole)
        box = item.data(0, Qt.ItemDataRole.UserRole + 1)
        group_shipment = item.data(0, Qt.ItemDataRole.UserRole + 2)
        
        if shipment and not box and not group_shipment:
            shipment.is_expanded = False
            # Schedule a delayed save to the database to avoid performance issues
            if hasattr(self.shipment_manager, 'schedule_save'):
                self.shipment_manager.schedule_save()
            
        elif group_shipment:
            group_shipment.is_expanded = False
            # Schedule a delayed save to the database to avoid performance issues
            if hasattr(self.shipment_manager, 'schedule_save'):
                self.shipment_manager.schedule_save()

    def show_shipment_properties(self, shipment_name):
        self.shipment_manager.show_shipment_properties(shipment_name)

    def show_group_shipment_properties(self, group_shipment):
        from dialogs import GroupShipmentPropertiesDialog
        dialog = GroupShipmentPropertiesDialog(group_shipment, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            for shipment in group_shipment.sub_shipments.values():
                self.data_controller.save_shipment(shipment)
            self.statusBar().showMessage(f"Свойства групповой поставки '{group_shipment.group_name}' сохранены", 3000)
    
    # Система блокировки ппоставккуавок удалена, функция exit_shipment большеше не нуженна
    # def exit_shipment(self, shipment):
    #     """Выйти из поставки - освободить блокировку и снять выделение"""
    #     try:
    #         # Освобождаем блокировку поставки
    #         self.release_shipment_lock(shipment)
    #         
    #         # Снимаем выделение с поставки если она текущующая
    #         if self.current_shipment == shipment:
    #             self.current_shipment = None
    #         
    #         # Обновляем интерфейс
    #         if self.ui_updater:
    #             self.ui_updater.update_shipments_tree()
    #         
    #         self.statusBar().showMessage(f"Вышли из поставки: {shipment.destination_name}", 3000)
    #         
    #     except Exception as e:
    #         self.logger.error(f"Ошибка при выходе из поставки: {e}", exc_info=True)

    def show_shipment_context_menu(self, position):
        try:
            item = self.shipments_tree_widget.itemAt(position)
            if not item:
                return
    
            menu = QMenu(self)
            # Устанавливаем стиль для уменьшения расстояния между иконкой и текстом
            menu.setStyleSheet("""
                QMenu::item {
                    padding: 6px 12px 6px 20px; /* Уменьшаем левый отступ для уменьшения расстояния до иконки */
                }
                QMenu::icon {
                    padding-right: 2px; /* Устанавливаем расстояние между иконкой и текстом */
                }
            """)
            
            # Получаем данные из пользовательских ролей
            shipment = item.data(0, Qt.ItemDataRole.UserRole)
            box = item.data(0, Qt.ItemDataRole.UserRole + 1)
            group_shipment = item.data(0, Qt.ItemDataRole.UserRole + 2)
            
            if shipment and not box and not group_shipment: # Это обычная поставка
                # Загружаем и уменьшаем иконки
                archive_icon_path = str(config.get_resource_path(Path("Res") / "archive.png"))
                archive_pixmap = QPixmap(archive_icon_path).scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                archive_icon = QIcon(archive_pixmap)
                
                excel_icon_path = str(config.get_resource_path(Path("Res") / "excel.png"))
                excel_pixmap = QPixmap(excel_icon_path).scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                excel_icon = QIcon(excel_pixmap)
                
                word_icon_path = str(config.get_resource_path(Path("Res") / "word.png"))
                word_pixmap = QPixmap(word_icon_path).scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                word_icon = QIcon(word_pixmap)
                
                delete_icon_path = str(config.get_resource_path(Path("Res") / "delete.png"))
                delete_pixmap = QPixmap(delete_icon_path).scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                delete_icon = QIcon(delete_pixmap)
                
                clean_icon_path = str(config.get_resource_path(Path("Res") / "clean.png"))
                clean_pixmap = QPixmap(clean_icon_path).scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                clean_icon = QIcon(clean_pixmap)
                
                refresh_icon_path = str(config.get_resource_path(Path("Res") / "refresh.png"))
                refresh_pixmap = QPixmap(refresh_icon_path).scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                refresh_icon = QIcon(refresh_pixmap)
                
                rename_icon_path = str(config.get_resource_path(Path("Res") / "rename.png"))
                rename_pixmap = QPixmap(rename_icon_path).scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                rename_icon = QIcon(rename_pixmap)
                
                settings_icon_path = str(config.get_resource_path(Path("Res") / "settings.png"))
                settings_pixmap = QPixmap(settings_icon_path).scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                settings_icon = QIcon(settings_pixmap)
                
                # Добавляем иконку для импорта
                import_icon_path = str(config.get_resource_path(Path("Res") / "import.png"))
                import_pixmap = QPixmap(import_icon_path).scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                import_icon = QIcon(import_pixmap)
                
                # Доба��ляем пункт архивации
                archive_action = menu.addAction(archive_icon, "Отправить в архив")
                archive_action.triggered.connect(lambda: self.archive_shipment(shipment.destination_name))
                
                export_action = menu.addAction(excel_icon, "Экспорт коробок")
                import_action = menu.addAction(import_icon, "Импорт коробок")
                
                # Добавля��м пункт "��ист на паллет"
                packing_list_action = menu.addAction(word_icon, "Лист на паллет")
                packing_list_action.triggered.connect(lambda: self.generate_packing_list())
                
                update_action = menu.addAction(refresh_icon, "Обновить состав поставки")
                rename_action = menu.addAction(rename_icon, "Переименовать поставки")
                delete_action = menu.addAction(delete_icon, "Удалить поставки")
                # Добавляем пункт "Удалить пустые коробки" для поставки
                delete_empty_boxes_action = menu.addAction(clean_icon, "Удалить пустые коробки")
                
                menu.addSeparator()
                
                properties_action = menu.addAction(settings_icon, "Свойства")
                
                # Система бл��кировки ппоставккуавок удалена, пункт "Выйти из поставки" больше��е не нужен��н
                
                action = menu.exec(self.shipments_tree_widget.mapToGlobal(position))
                if action == export_action: # Обработкатка экспорта коробкуок
                    # Устанавливаем выбранную ппоставкку��вку как текущующую
                    self.current_shipment = shipment
                    self.export_boxes()
                elif action == import_action: # Обработкатка импорта коробкуок
                    # Устанавливаем выбранную ппоставккуавкиу как текущующую
                    self.current_shipment = shipment
                    self.shipment_manager.import_boxes()
                elif action == packing_list_action:  # Обработка��ка "Лист на паллет"
                    # Устанавливаем выбранную ппоставккуавкиу как текущующую
                    self.current_shipment = shipment
                    self.generate_packing_list()
                elif action == update_action:
                    # Устанавливаем выбранную ��поставкку как текущую��ую
                    self.current_shipment = shipment
                    self.shipment_operations.update_shipment_composition()
                elif action == rename_action:
                    # Устанавливаем выбранную ппоставккуавкиу ��как текущующую
                    self.current_shipment = shipment
                    self.shipment_operations.rename_shipment(shipment.destination_name)
                elif action == delete_action:
                    # Устанавливаем выбранную ппоставккуавкиу как текущующую
                    self.current_shipment = shipment
                    self.shipment_operations.delete_shipment(shipment.destination_name)
                elif action == delete_empty_boxes_action: # Обработкатка удаления пустых коробкуок
                    # Устанавливаем выбранную ппоставккуавкиу как текущующую
                    self.current_shipment = shipment
                    self.shipment_manager.delete_empty_boxes(shipment)
                elif action == properties_action:
                    # Устанавливаем выбранную ппоставккуавкиу как текущующую
                    self.current_shipment = shipment
                    self.show_shipment_properties(shipment.destination_name)
# elif 'exit_action' in locals() and action == exit_action:  # Обработкатка выхода из поставки
                #     # Устанавливаем выбранную ппоставккуавкиу как текущующую
                #     self.current_shipment = shipment
                #     self.exit_shipment(shipment)

            elif group_shipment:  # Это группуовая поставка
                archive_action = menu.addAction("📦 Отправить группуу в архив")
                archive_action.triggered.connect(lambda: self.archive_group_shipment(group_shipment.group_name))
                menu.addSeparator()

                update_group_action = menu.addAction("Обновить состав групповой поставки")
                update_gsheets_action = menu.addAction("Обновить из Google Sheets")
                export_all_action = menu.addAction("Экспорт всех коробок")
                rename_group_action = menu.addAction("Переименовать группу")
                delete_group_action = menu.addAction("Удалить группу")

                menu.addSeparator()
                settings_icon_path = str(config.get_resource_path(Path("Res") / "settings.png"))
                settings_pixmap = QPixmap(settings_icon_path).scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                settings_icon = QIcon(settings_pixmap)
                properties_group_action = menu.addAction(settings_icon, "Свойства")

                action = self.shipments_tree_widget.mapToGlobal(position)
                action = menu.exec(action)
                if action == update_group_action:
                    self.shipment_operations.update_group_shipment_composition(group_shipment)
                elif action == update_gsheets_action:
                    self.update_group_shipment_from_google_sheets(group_shipment)
                elif action == export_all_action:
                    self.shipment_manager.export_all_group_boxes(group_shipment)
                elif action == rename_group_action:
                    self.shipment_operations.rename_group_shipment(group_shipment.group_name)
                elif action == delete_group_action:
                    self.shipment_operations.delete_group_shipment(group_shipment.group_name)
                elif action == properties_group_action:
                    self.show_group_shipment_properties(group_shipment)
            
            elif shipment and box and not group_shipment:  # Это коробкука
                # Добавляем пункт печати этикетки в начало меню
                print_label_action = menu.addAction("Печать этикетки")
                menu.addSeparator()
                
                delete_action = menu.addAction("Удалить коробку")
                rename_action = menu.addAction("Переименовать коробку")
                
                action = menu.exec(self.shipments_tree_widget.mapToGlobal(position))
                if action == print_label_action:
                    self.print_label(shipment, box)
                elif action == delete_action:
                    # Находим индекс коробки для удаления
                    index = -1
                    for i, b in enumerate(shipment.boxes):
                        if b == box:
                            index = i
                            break
                    if index >= 0:
                        self.shipment_manager.delete_box(index)
                elif action == rename_action:
                    # Находим индекс коробки для переименования��
                    index = -1
                    for i, b in enumerate(shipment.boxes):
                        if b == box:
                            index = i
                            break
                    if index >= 0:
                        self.shipment_manager.rename_box(index)
        except Exception as e:
            self.logger.error(f"Ошибка при отображении контекстного меню поставки: {e}", exc_info=True)

    def show_box_table_context_menu(self, position):
        self.shipment_manager.show_box_table_context_menu(position)

    def delete_shipment(self, shipment_name):
        self.shipment_operations.delete_shipment(shipment_name)

    def rename_shipment(self, old_name):
        self.shipment_operations.rename_shipment(old_name)

    def delete_box(self, index):
        self.shipment_manager.delete_box(index)

    def rename_box(self, index):
        self.shipment_manager.rename_box(index)

    def remove_item_from_box(self, barcode):
        self.shipment_manager.remove_item_from_box(barcode)

    def export_boxes(self):
        self.shipment_manager.export_boxes()

    def print_label(self, shipment=None, box=None):
        """Печать этикетки для коробку"""
        try:
            if shipment and box:
                # Используем display_name для этикеток (без преф��кса группуы)
                destination_name = getattr(shipment, 'display_name', shipment.destination_name)
                
                # П��ямая печать транспортной этикетки на коробкуку (направление и номер)
                from label_print_dialog import create_and_print_transport_label
                from PyQt6.QtPrintSupport import QPrinter
                printer = QPrinter()
                create_and_print_transport_label(printer, destination_name, box.box_id)
                return
            
            # Если контекст не передан, открываем общий диалог печати
            from label_print_dialog import LabelPrintDialog
            dialog = LabelPrintDialog(self)
            dialog.exec()
        except ImportError:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "Ошибка", "Модуль печати этикеток не найден. Убедитесь, что файл label_print_dialog.py находится в проекте.")
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Ошибка при печати этикетки: {e}", exc_info=True)
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Ошибка", f"Ошибка при печати этикетки:\n{e}")

    def load_sku_data_from_db(self):
        """Загружает данные из общей базы данных в память"""
        sku_data = {}
        try:
            # Используем общую базу данных PostgreSQL
            from database import execute_query
            rows = execute_query("SELECT barcode, article, name FROM sku", fetchall=True)
            
            # Загружаем данные из базы
            for row in rows:
                barcode, article, name = row
                sku_data[barcode] = {'article': article, 'name': name}
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Не удалось загрузить данные из базы: {e}", exc_info=True)
        
        return sku_data

    def generate_packing_list(self):
        """Создать лист на паллет для текующей поставки"""
        try:
            if not self.current_shipment:
                QMessageBox.warning(self, "Ошибка", "Сначала выберите поставку!")
                return

            if not DOCX_AVAILABLE:
                QMessageBox.warning(self, "Ошибка",
                                  "Для создания листа на паллет необходимо установить библиотекущую python-docx!\n\n"
                                  "Установите ее командой: pip install python-docx")
                return

            # Получаем свойства поставки
            props = self.current_shipment.properties

            # Проверяем обязательные поля
            if not props.shipment_number:
                QMessageBox.warning(self, "Ошибка", "В свойствах поставки не указан номер поставки!")
                return

            # Запрашиваем путь для сохранения
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Сохранить лист на паллет",
                f"Лист на паллет {props.shipment_number}.docx",
                "Word Documents (*.docx)"
            )

            if not file_path:
                return

            try:
                # Создаем новый документ
                doc = Document()

                # Устанавливаем размер страницы и поля
                section = doc.sections[0]
                section.page_height = Inches(11)
                section.page_width = Inches(8.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)

                # Добавляем заголовок
                title = doc.add_paragraph('Лист на паллет.')
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].bold = True
                title.runs[0].font.size = Pt(24)  # Увеличиваем размер шрифта

                # Добавляем пустую строку
                doc.add_paragraph()

                # Создаем таблицу для данных
                table = doc.add_table(rows=9, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'

                # Заполняем таблицу данными
                data = [
                    ("Номер паллеты", "1"),
                    ("Кол-во паллет", "1"),
                    ("Номер поставки", props.shipment_number or ""),
                    ("Коробов на паллете", str(len(self.current_shipment.boxes))),
                    ("Коробов в поставке", str(len(self.current_shipment.boxes))),
                    ("Склад", props.destination_warehouse or ""),
                    ("Способ поставки", "Короб"),
                    ("Поставщик, ИНН", self.get_supplier_info(props.legal_entity)),
                    ("Дата поставки", self.format_shipment_date(props.shipment_date))
                ]

                for i, (key, value) in enumerate(data):
                    row = table.rows[i]
                    
                    # Левая колонка - название поля
                    left_cell = row.cells[0]
                    left_cell.text = key
                    left_paragraph = left_cell.paragraphs[0]
                    left_paragraph.runs[0].bold = True
                    left_paragraph.runs[0].font.size = Pt(24)  # Увеличиваем размер шрифта
                    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Правая колонка - значение
                    right_cell = row.cells[1]
                    right_cell.text = value
                    right_paragraph = right_cell.paragraphs[0]
                    right_paragraph.runs[0].font.size = Pt(24)  # Увеличиваем размер шрифта
                    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Выделяем жирным важные поля
                    if key in ["Номер поставки", "Коробов на паллете"]:
                        right_paragraph.runs[0].bold = True

                # Настраиваем ширину колонок и высоту строк
                for row in table.rows:
                    # Увеличиваем ширину колонок
                    row.cells[0].width = Inches(3.5)
                    row.cells[1].width = Inches(4.0)
                    
                    # Увеличиваем высоту строк
                    for cell in row.cells:
                        cell.paragraphs[0].paragraph_format.space_after = Pt(12)
                        cell.paragraphs[0].paragraph_format.space_before = Pt(12)

                # Добавляем отступ после таблицы
                doc.add_paragraph()

                # Сохраняем документ
                doc.save(file_path)

                utils.play_sound(self.ok_sound, self.tone_sound)
                # Статус скрыт, соо��щения не отображаются
                # self.statusBar().showMessage(f"Лист на паллет сохранен: {file_path}", 5000)
                QMessageBox.information(self, "Успех", "Лист на паллет успешно создан!")

            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось создать лист на паллет:\n{str(e)}")
        except Exception as e:
            self.logger.error(f"Ошибка при создании листа на паллет: {e}", exc_info=True)
            QMessageBox.critical(self, "Ошибка", f"Ошибка при создании листа на паллет:\n{e}")
        
    def get_supplier_info(self, legal_entity):
        """Получить информацию о ппоставккуавщике по выбранному юр. лицу"""
        if legal_entity == "ООО ОНДЕФОР":
            return 'ООО "ОНДЕФОР ГРУПП" 5029279234'
        elif legal_entity == "ИП Лазарчук":
            return "ИП Лазарчук К.Е., 632410867452"
        else:
            return legal_entity or ""

    def format_shipment_date(self, date_str):
        """Форматировать дату поставки"""
        if not date_str:
            from datetime import datetime
            return datetime.now().strftime("%d.%m.%Y")
        
        try:
            from datetime import datetime
            # Пытаемся разобрать дату в формате "yyyy-MM-dd"
            date_obj = datetime.strptime(date_str, "%Y-%m-%d")
            return date_obj.strftime("%d.%m.%Y")
        except ValueError:
            # Если не удалось разобрать, возвращаем как есть
            return date_str

    def showEvent(self, event):
        super().showEvent(event)
        # Восстанавливаем размеры сплиттера при показе окна
        QTimer.singleShot(100, self.load_window_state)
        
        # Обновляем таблицы для обеспечения отображения данных
        if hasattr(self, 'shipment_table') and self.shipment_table:
            QTimer.singleShot(250, lambda: self.shipment_table.updateGeometry())
        if hasattr(self, 'current_box_table') and self.current_box_table:
            QTimer.singleShot(250, lambda: self.current_box_table.updateGeometry())
            
        # Дополнительно вызываем update_ui через 300 мс для гарантии отображения всех данных
        QTimer.singleShot(300, self.update_ui)
        
        # Обновляем геометрию сплиттера через 400 мс для гарантии правильной перерисовки
        if hasattr(self, 'main_splitter'):
            QTimer.singleShot(400, lambda: self.main_splitter.updateGeometry())

    def open_archive(self):
        """Открыть окно архива"""
        try:
            # Создаем новый экземпляр диалога архива каждый раз
            archive_dialog = ArchiveDialog(self)
            # Устанавливаем флаг, чтобы диалог уничтожался после закрытия
            archive_dialog.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose, True)
            # Показываем диалог как модальное окно
            archive_dialog.exec()
        except Exception as e:
            self.logger.error(f"Ошибка при открытии архива: {e}", exc_info=True)
            QMessageBox.critical(self, "Ошибка", f"Ошибка при открытии архива:\n{e}")

    def archive_shipment(self, shipment_name):
        """Архивировать поставку (асинхронно)"""
        try:
            if not shipment_name:
                return
                
            reply = QMessageBox.question(
                self, "Подтверждение архивации",
                f"Вы уверены, что хотите отправить поставку «{shipment_name}» в архив?\n\nПоставка будет скрыта из основного списка, но сохранена в базе данных.",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                self.force_save_session()
                self.show_progress(f"Архивация поставки «{shipment_name}»...")
                
                self.async_manager.execute_async(
                    self.data_controller.archive_shipment,
                    lambda result: self._on_archive_shipment_finished(result, shipment_name),
                    lambda err: self._on_archive_shipment_error(err, shipment_name),
                    shipment_name, self.current_user
                )
        except Exception as e:
            self.logger.error(f"Ошибка при архивации поставки {shipment_name}: {e}", exc_info=True)
            QMessageBox.critical(self, "Ошибка", f"Ошибка при архивации поставки:\n{e}")

    def _on_archive_shipment_finished(self, success, shipment_name):
        """Обработка успешной архивации поставки"""
        if success:
            self.load_all_data()
            self.hide_progress(f"Поставка «{shipment_name}» архивирована", 2000)
        else:
            self.hide_progress("Ошибка архивации", 3000)
            QMessageBox.critical(self, "Ошибка", "Не удалось отправить поставку в архив")

    def _on_archive_shipment_error(self, error_msg, shipment_name):
        """Обработка ошибки архивации поставки"""
        self.logger.error(f"Ошибка при архивации поставки {shipment_name}: {error_msg}")
        self.hide_progress("Ошибка архивации", 3000)
        QMessageBox.critical(self, "Ошибка", f"Ошибка при архивации поставки:\n{error_msg}")

    def archive_group_shipment(self, group_name):
        """Архивировать групповую поставку (асинхронно)"""
        try:
            if not group_name:
                return
                
            reply = QMessageBox.question(
                self, "Подтверждение архивации",
                f"Вы уверены, что хотите отправить групповую поставку «{group_name}» в архив?\n\nВсе входящие в нее поставки также будут архивированы.",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                self.force_save_session()
                self.show_status(f"Архивация групповой поставки «{group_name}»...", 3000)
                
                # Собираем имена поставок для архивации
                group_shipment = self.group_shipments[group_name]
                shipment_names = []
                for shipment_name, shipment in group_shipment.sub_shipments.items():
                    if hasattr(shipment, 'original_destination_name'):
                        shipment_names.append(shipment.original_destination_name)
                    else:
                        shipment_names.append(shipment.destination_name)
                
                self.async_manager.execute_async(
                    self._archive_group_shipments_batch,
                    callback=lambda result: self._on_archive_group_shipment_finished(result, group_name),
                    error_callback=lambda err: self._on_archive_group_shipment_error(err, group_name),
                    shipment_names=shipment_names
                )
        except Exception as e:
            self.logger.error(f"Ошибка при архивации групповой поставки {group_name}: {e}", exc_info=True)
            QMessageBox.critical(self, "Ошибка", f"Ошибка при архивации групповой поставки:\n{e}")

    def _archive_group_shipments_batch(self, shipment_names):
        """Архивирует все подпоставки в фоне"""
        for name in shipment_names:
            self.data_controller.archive_shipment(name, self.current_user)
        return True

    def _on_archive_group_shipment_finished(self, success, group_name):
        """Обработка успешной архивации групповой поставки"""
        if success:
            self.load_all_data()
            self.hide_progress(f"Групповая поставка «{group_name}» архивирована", 2000)
        else:
            self.hide_progress("Ошибка архивации", 3000)
            QMessageBox.critical(self, "Ошибка", "Не удалось отправить групповую поставку в архив")

    def _on_archive_group_shipment_error(self, error_msg, group_name):
        """Обработка ошибки архивации групповой поставки"""
        self.logger.error(f"Ошибка при архивации групповой поставки {group_name}: {error_msg}")
        self.hide_progress("Ошибка архивации", 3000)
        QMessageBox.critical(self, "Ошибка", f"Ошибка при архивации групповой поставки:\n{error_msg}")

    def start_shipment_check(self):
        """Начать проверку поставки с импорта Excel файла"""
        try:
            # Create new shipment check dialog each time to avoid state issues
            shipment_check_dialog = ShipmentCheckDialog(self)
            # Show the dialog
            if shipment_check_dialog.exec() == QDialog.DialogCode.Accepted:
                # Статус скрыт, сообщения не отображаются
                # self.statusBar().showMessage("Проверка поставки начата", 3000)
                pass
            
        except Exception as e:
            self.logger.error(f"Ошибка при запуске проверки поставки: {e}", exc_info=True)
            QMessageBox.critical(self, "Ошибка", f"Не удалось запустить проверку поставки:\n{e}")
    
    def show_about(self):
        """Показать информацию о программе"""
        try:
            # Получаем версию
            try:
                from version import get_version_string
                version = get_version_string()
            except Exception:
                version = "1.0.0.0"
            
            about_text = (
                f"<h2>WB Packer</h2>"
                f"<p><b>Система сборки поставок</b></p>"
                f"<p>Версия: <b>{version}</b></p>"
                f"<p>© 2025 Все права защищены</p>"
                f"<p><b>Василий Булеков</b></p>"
            )
            QMessageBox.about(self, "О программе", about_text)
        except Exception as e:
            self.logger.error(f"Ошибка при отображении информации о программе: {e}", exc_info=True)
    
    def update_moysklad_button_visibility(self):
        """Обновляет видимость кнопки синхронизации с МойСклад в зависимости от настроек"""
        if hasattr(self, 'moysklad_sync_btn') and self.moysklad_sync_btn:
            try:
                # Настройки МойСклад теперь глобальные
                is_enabled = database.get_moysklad_enabled()
                self.moysklad_sync_btn.setVisible(bool(is_enabled))
            except Exception as e:
                self.logger.error(f"Ошибка при обновлении видимости кнопки МойСклад: {e}", exc_info=True)
                # В случае ошибки, показываем кнопку
                self.moysklad_sync_btn.setVisible(True)

    def sync_moysklad_stocks(self):
        """
        Синхронизация остатков с МойСклад: обновляет значения текущующих остатков на складах (которые были заданы в настройках) для всех неархивных поставок
        """
        try:
            import database
            from PyQt6.QtWidgets import QMessageBox
            
            self.logger.info("Начало улучшенной синхронизации остатков с МойСклад для всех неархивных поставок")
            self.show_status("Синхронизация остатков с МойСклад...")

            # Проверяем, включена ли интеграция
            if not database.get_moysklad_enabled():
                self.logger.warning("Интеграция с МойСклад отключена")
                self.show_status("Интеграция с МойСклад отключена", 3000)
                QMessageBox.warning(self, "Ошибка", "Интеграция с МойСклад отключена. Включите интеграцию в настройках.")
                return

            # Проверяем наличие токена МойСклад (теперь глобальный)
            if not database.get_moysklad_token():
                self.logger.warning("Токен МойСклад не настроен")
                self.show_status("Токен МойСклад не настроен", 3000)
                QMessageBox.warning(self, "Ошибка", "Токен МойСклад не настроен. Проверьте настройки интеграции.")
                return

            self.logger.info("Токен МойСклад настроен")
            
            # Получаем все неархивные поставки
            all_shipments = {}
            all_shipments.update(self.shipments)
            for group_shipment in self.group_shipments.values():
                all_shipments.update(group_shipment.sub_shipments)
            
            # Фильтруем только неархивные поставки
            non_archived_shipments = {}
            for name, shipment in all_shipments.items():
                if not shipment.archived:
                    non_archived_shipments[name] = shipment
            
            self.logger.info(f"Найдено {len(non_archived_shipments)} неархивных поставок для синхронизации остатков")
            
            if not non_archived_shipments:
                self.logger.info("Нет неархивных поставок для синхронизации остатков")
                QMessageBox.information(self, "Информация", "Нет неархивных поставок для синхронизации остатков.")
                # Статус скрыт, сообщения не отображаются
                # self.statusBar().showMessage("Синхронизация остатков завершена", 3000)
                return
            
            # Запускаем асинхронную синхронизацию
            from improved_moysklad_sync import ImprovedMoyskladSync
            self.improved_sync_handler = ImprovedMoyskladSync(self.ui_updater)
            
            # Подключаем сигналы для обновления прогресса
            self.improved_sync_handler.progress_updated.connect(
                lambda current, total: self.update_progress(current, f"Синхронизация МойСклад... ({current}/{total})")
            )
            self.improved_sync_handler.sync_completed.connect(self._on_sync_completed)
            self.improved_sync_handler.sync_error.connect(self._on_sync_error)
            self.improved_sync_handler.sync_started.connect(
                lambda: self.show_progress("Начало синхронизации остатков с МойСклад...")
            )
            
            # Запускаем асинхронную синхронизацию
            self.improved_sync_handler.sync_stocks_async(non_archived_shipments)
            
            self.logger.info("Асинхронная синхронизация запущена")
            
        except ImportError as e:
            self.logger.error(f"Модуль интеграции с МойСклад не найден: {e}")
            QMessageBox.warning(self, "Ошибка", f"Модуль интеграции с МойСклад не найден. Убедитесь, что файлы moysklad_api.py и improved_moysklad_sync.py находятся в проекте.\n{e}")
        except Exception as e:
            self.logger.error(f"Ошибка при запуске синхронизации остатков с МойСклад: {e}", exc_info=True)
            QMessageBox.critical(self, "Ошибка", f"Ошибка при запуске синхронизации остатков с МойСклад:\n{e}")
    
    def _on_sync_completed(self, result):
        """Обработкатка завершения асинхронной синхронизации"""
        try:
            import database
            from PyQt6.QtWidgets import QApplication
            
            # Обновляем кэш с новыми данными в базе данных
            self.logger.info("Обновление кэша остатков в базе данных")
            for barcode, quantity in result.items():
                database.set_stock_cache(barcode, quantity)
            
            # Если есть текущующая поставка, обновляем интерфейс для нее
            if self.current_shipment:
                self.logger.info("Принудительное обновление интерфейса таблицы текущей поставки")

                vertical_scroll_pos = self.shipment_table.verticalScrollBar().value()
                horizontal_scroll_pos = self.shipment_table.horizontalScrollBar().value()

                self.ui_updater.update_shipment_table()

                self.shipment_table.verticalScrollBar().setValue(vertical_scroll_pos)
                self.shipment_table.horizontalScrollBar().setValue(horizontal_scroll_pos)
                self.restore_table_columns_width()
                self.shipment_table.updateGeometry()
            elif self.current_shipment is None:
                for group_name, group_shipment in self.group_shipments.items():
                    if self.shipments_tree_widget.currentItem():
                        current_item = self.shipments_tree_widget.currentItem()
                        if current_item.data(0, Qt.ItemDataRole.UserRole + 2) == group_shipment:
                            self.ui_updater.update_group_shipment_summary(group_shipment)
                            self.ui_updater.update_group_shipment_boxes_table(group_shipment)
                            self.ui_updater.update_group_shipment_items_table(group_shipment)
                            break
            
            # Обновляем статус
            all_shipments = {}
            all_shipments.update(self.shipments)
            for group_shipment in self.group_shipments.values():
                all_shipments.update(group_shipment.sub_shipments)
            
            non_archived_shipments = {name: shipment for name, shipment in all_shipments.items() if not shipment.archived}
            
            self.hide_progress(f"Синхронизация завершена: {len(non_archived_shipments)} поставок", 3000)
            self.logger.info("Синхронизация остатков завершена успешно")
        except Exception as e:
            self.logger.error(f"Ошибка при обработке завершения синхронизации: {e}", exc_info=True)
            self.hide_progress("Ошибка синхронизации", 3000)
    
    def _on_sync_error(self, error_message):
        """Обработка ошибки асинхронной синхронизации"""
        from PyQt6.QtWidgets import QMessageBox
        self.logger.error(f"Ошибка при синхронизации остатков с МойСклад: {error_message}")
        self.hide_progress("Ошибка синхронизации", 3000)
        QMessageBox.critical(self, "Ошибка", f"Ошибка при синхронизации остатков с МойСклад:\n{error_message}")

    def open_check_stock_dialog(self):
        """
        Открытие диалога проверки остатков по штрихкоду
        """
        try:
            from check_stock_dialog import CheckStockDialog
            dialog = CheckStockDialog(self)
            dialog.exec()
        except ImportError:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "Ошибка", "Модуль проверки остатков не найден. Убедитесь, что файл check_stock_dialog.py находится в проекте.")
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Ошибка при открытии диалога проверки остатков: {e}", exc_info=True)
            QMessageBox.critical(self, "Ошибка", f"Ошибка при открытии диалога проверки остатков:\n{e}")

    def open_label_print_dialog(self):
       """Открытие диалога печати этикеток"""
       try:
           from label_print_dialog import LabelPrintDialog
           dialog = LabelPrintDialog(self)
           dialog.exec()
       except ImportError:
           from PyQt6.QtWidgets import QMessageBox
           QMessageBox.warning(self, "Ошибка", "Модуль печати этикеток не найден. Убедитесь, что файл label_print_dialog.py находится в проекте.")
       except Exception as e:
           import logging
           logger = logging.getLogger(__name__)
           logger.error(f"Ошибка при открытии диалога печати этикеток: {e}", exc_info=True)

    def update_sku_table_async(self):
        """Асинхронное обновление таблицы SKU из Google Sheets"""
        self.show_busy_progress("Обновление таблицы SKU...")
        self.logger.info("Запуск асинхронного обновления таблицы SKU из Google Sheets")
        
        self.async_manager.execute_async(
            self._update_sku_table_worker,
            callback=self._on_sku_update_finished,
            error_callback=self._on_sku_update_error
        )

    def _update_sku_table_worker(self):
        """Worker для обновления SKU в фоновом потоке из Google Sheets"""
        import os
        import time
        from db_connection import get_connection, get_db_type
        
        t0 = time.time()
        
        try:
            import gspread
            from google.oauth2.service_account import Credentials
        except ImportError:
            raise ImportError("Установите gspread: pip install gspread google-auth")
        
        # Загрузка credentials из файла service account
        # При работе из EXE файл извлекается из встроенных ресурсов через _MEIPASS
        if getattr(sys, 'frozen', False):
            # В EXE: файл распаковывается во временную папку _MEIPASS
            credentials_path = os.path.join(sys._MEIPASS, "e-object-470910-p6-3500f3ddbdd3.json")
        else:
            # При разработке: файл рядом с исходным кодом
            credentials_path = os.path.join(os.path.dirname(__file__), "e-object-470910-p6-3500f3ddbdd3.json")
        if not os.path.exists(credentials_path):
            raise FileNotFoundError(f"Файл credentials не найден: {credentials_path}")
        
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds = Credentials.from_service_account_file(credentials_path, scopes=scope)
        client = gspread.authorize(creds)
        
        self.logger.info(f"[SKU] Авторизация: {time.time()-t0:.2f}s")
        
        # ID таблицы из URL
        spreadsheet_id = "1tQzh_qTnldbpeu9ryNF8ZKY4-amwT8UfuMqbSU1qOlA"
        spreadsheet = client.open_by_key(spreadsheet_id)
        sheet = spreadsheet.sheet1
        
        # Используем get_all_values() вместо get_all_records() — в 3-5 раз быстрее
        # get_all_records() создаёт dict для каждой строки, что очень медленно на больших таблицах
        all_values = sheet.get_all_values()
        
        self.logger.info(f"[SKU] Получено {len(all_values)} строк из Google Sheets: {time.time()-t0:.2f}s")
        
        if len(all_values) < 2:
            raise ValueError("Таблица SKU пуста или не удалось загрузить данные")
        
        # Находим индексы колонок по заголовкам
        headers = [h.strip().lower() for h in all_values[0]]
        
        # Ищем колонку штрихкода
        barcode_idx = None
        for i, h in enumerate(headers):
            if h in ('штрихкод', 'barcode', 'шк', 'штрих-код', 'баркод'):
                barcode_idx = i
                break
        if barcode_idx is None:
            # Если не нашли по заголовку, берём первую колонку
            barcode_idx = 0
        
        # Ищем колонку артикула
        article_idx = None
        for i, h in enumerate(headers):
            if h in ('артикул', 'article', 'арт', 'sku'):
                article_idx = i
                break
        
        # Ищем колонку наименования
        name_idx = None
        for i, h in enumerate(headers):
            if h in ('наименование', 'name', 'название', 'товар'):
                name_idx = i
                break
        
        self.logger.info(f"[SKU] Колонки: barcode={barcode_idx}, article={article_idx}, name={name_idx}")
        
        # Парсим данные напрямую из списка списков
        records_to_insert = []
        records_with_labels = 0
        skipped_empty_barcode = 0
        seen_barcodes = {}  # Для дедупликации: barcode -> (article, name)
        
        for row in all_values[1:]:  # Пропускаем заголовок
            if not row or len(row) <= barcode_idx:
                skipped_empty_barcode += 1
                continue
            
            barcode = row[barcode_idx].strip() if barcode_idx < len(row) else ""
            article = row[article_idx].strip() if article_idx is not None and article_idx < len(row) else ""
            name = row[name_idx].strip() if name_idx is not None and name_idx < len(row) else ""
            
            # Нормализуем штрихкод
            barcode = barcode.replace(" ", "").replace("-", "").replace("\t", "")
            article = article.replace(" ", "").replace("\t", "")
            
            if not barcode or barcode.lower() in ('nan', 'none', ''):
                skipped_empty_barcode += 1
                continue
            
            # Дедупликация: последняя запись побеждает
            seen_barcodes[barcode] = (article if article else None, name if name else None)
        
        # Преобразуем в список для вставки
        for barcode, (article, name) in seen_barcodes.items():
            if name and name.lower() not in ('nan', 'none', ''):
                records_with_labels += 1
            records_to_insert.append((barcode, article, name))
        
        self.logger.info(f"[SKU] Подготовлено {len(records_to_insert)} записей ({records_with_labels} с именами, дедупликация из {len(all_values)-1} строк): {time.time()-t0:.2f}s")
        
        if not records_to_insert:
            raise ValueError("Нет валидных записей для обновления SKU")
        
        # Массовая вставка
        db_type = get_db_type()
        conn = get_connection()
        cursor = conn.cursor()
        
        try:
            if db_type == "sqlite":
                cursor.execute("DELETE FROM sku")
                cursor.executemany(
                    "INSERT INTO sku (barcode, article, name) VALUES (?, ?, ?) ON CONFLICT(barcode) DO UPDATE SET article=excluded.article, name=excluded.name",
                    records_to_insert
                )
            else:
                from db_connection import psycopg, psycopg2
                cursor.execute("DELETE FROM sku")
                if psycopg is not None:
                    cursor.executemany(
                        "INSERT INTO sku (barcode, article, name) VALUES (%s, %s, %s) ON CONFLICT (barcode) DO UPDATE SET article = EXCLUDED.article, name = EXCLUDED.name",
                        records_to_insert
                    )
                else:
                    import psycopg2.extras
                    psycopg2.extras.execute_values(
                        cursor,
                        "INSERT INTO sku (barcode, article, name) VALUES %s ON CONFLICT (barcode) DO UPDATE SET article = EXCLUDED.article, name = EXCLUDED.name",
                        records_to_insert,
                        page_size=1000
                    )
            
            conn.commit()
            self.logger.info(f"[SKU] Вставка завершена: {time.time()-t0:.2f}s")
        except Exception:
            conn.rollback()
            raise
        finally:
            cursor.close()
        
        # Сбрасываем кэш наименований чтобы UI подхватил новые данные
        from database import clear_product_names_cache
        clear_product_names_cache()
        
        elapsed = round(time.time() - t0, 1)
        return {
            'records_with_labels': records_with_labels,
            'total_records': len(records_to_insert),
            'skipped': skipped_empty_barcode,
            'elapsed': elapsed,
        }

    def _on_sku_update_finished(self, result):
        """Обработка успешного обновления SKU"""
        msg = f"SKU обновлено: {result['records_with_labels']} с наименованиями из {result['total_records']} (пропущено: {result.get('skipped', 0)}) за {result.get('elapsed', '?')}с"
        self.hide_progress(msg, 5000)
        self.logger.info(msg)
        # Обновляем UI чтобы подхватить новые имена из SKU
        if self.current_shipment and self.ui_updater:
            self.logger.info("Обновляем таблицу поставки после SKU update")
            self.ui_updater.update_shipment_table()

    def _on_sku_update_error(self, error_msg):
        """Обработка ошибки обновления SKU"""
        self.hide_progress("Ошибка обновления SKU", 3000)
        self.logger.error(f"Ошибка обновления таблицы SKU: {error_msg}")
        from PyQt6.QtWidgets import QMessageBox
        QMessageBox.critical(self, "Ошибка", f"Ошибка обновления таблицы SKU:\n{error_msg}")

    def print_product_label(self):
       """Печать этикетки на товар (открывает диалог печати этикеток)"""
       try:
           from label_print_dialog import LabelPrintDialog
           dialog = LabelPrintDialog(self)
           dialog.exec()
       except ImportError:
           from PyQt6.QtWidgets import QMessageBox
           QMessageBox.warning(self, "Ошибка", "Модуль печати этикеток не найден. Убедитесь, что файл label_print_dialog.py находится в проекте.")
       except Exception as e:
           import logging
           logger = logging.getLogger(__name__)
           logger.error(f"Ошибка при открытии диалога печати этикетки на товар: {e}", exc_info=True)
           QMessageBox.critical(self, "Ошибка", f"Ошибка при открытии диалога печати этикетки на товар:\n{e}")

    def print_box_label(self):
       """Печать этикетки на коробкуку (открывает диалог печати этикеток)"""
       try:
           from label_print_dialog import LabelPrintDialog
           dialog = LabelPrintDialog(self)
           dialog.exec()
       except ImportError:
           from PyQt6.QtWidgets import QMessageBox
           QMessageBox.warning(self, "Ошибка", "Модуль печати этикеток не найден. Убедитесь, что файл label_print_dialog.py находится в проекте.")
       except Exception as e:
           import logging
           logger = logging.getLogger(__name__)
           logger.error(f"Ошибка при открытии диалога печати этикетки на коробкуку: {e}", exc_info=True)
           QMessageBox.critical(self, "Ошибка", f"Ошибка при открытии диалога печати этикетки на коробкуку:\n{e}")

    def eventFilter(self, obj, event):
        """
        Фильтр событий для перехвата нажатия пробела и перевода фокуса на поле ввода штрихкода
        """
        if event.type() == event.Type.KeyPress:
            if event.key() == Qt.Key.Key_Space.value:
                # Проверяем, находится ли фокус в каком-либо текстовом поле ввода
                focused_widget = self.focusWidget()
                
                # Если фокус уже в поле ввода штрихкода, не переносим фокус
                if focused_widget == self.scan_input:
                    return False  # Продолжаем стандартную обработку
                
                # Устанавливаем фокус на scan_input
                if hasattr(self, 'scan_input') and self.scan_input:
                    self.scan_input.setFocus()
                    self.scan_input.selectAll()  # Выделяем весь текст для удобства
                
                return True  # Событие обработано
        
        # Для других событий используем стандартную обработку
        return super().eventFilter(obj, event)
    
    def focus_scan_input_field(self):
        """
        Метод для перевода фокуса на поле ввода штрихкода
        """
        if hasattr(self, 'scan_input') and self.scan_input:
            self.scan_input.setFocus()
            self.scan_input.selectAll()  # Выделяем весь текст для удобства

    def show_status(self, message: str, timeout: int = 3000):
        """Отображает сообщение в статус-баре"""
        self.statusBar().showMessage(message, timeout)

    def show_progress(self, message: str, max_value: int = 100):
        """Показывает прогресс-бар в статус-баре"""
        self.status_progress_bar.setRange(0, max_value)
        self.status_progress_bar.setValue(0)
        self.status_progress_bar.setVisible(True)
        self.statusBar().showMessage(message)

    def show_busy_progress(self, message: str):
        """Показывает прогресс-бар в режиме занятости (бесконечная анимация)"""
        self.status_progress_bar.setRange(0, 0)
        self.status_progress_bar.setVisible(True)
        self.statusBar().showMessage(message)

    def update_progress(self, value: int, message: str = ""):
        """Обновляет прогресс-бар"""
        self.status_progress_bar.setValue(value)
        if message:
            self.statusBar().showMessage(message)

    def hide_progress(self, message: str = "Готово", timeout: int = 2000):
        """Скрывает прогресс-бар и показывает сообщение"""
        self.status_progress_bar.setRange(0, 100)
        self.status_progress_bar.setVisible(False)
        self.statusBar().showMessage(message, timeout)
